"""
DOCX Exporter — Consensus Pipeline v4.4

Converts Markdown content to Word (.docx), supporting mixed Chinese-English text.
Replaces PDF exporter (FPDF font compatibility issues in v2.5.3+).
Requires: python-docx
"""
import os
import re
from typing import Optional, List

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def markdown_to_docx(
    markdown_content: str,
    output_path: str,
    title: str = "Consensus Pipeline Report",
) -> str:
    """
    Convert Markdown content to Word (.docx) document.

    Args:
        markdown_content: Markdown-format content
        output_path: Output .docx file path
        title: Document title

    Returns:
        docx file path
    """
    if not HAS_DOCX:
        raise ImportError(
            "python-docx is required. Install with: pip install python-docx"
        )

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = markdown_content.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code block handling
        if line.strip().startswith("```"):
            if in_code_block:
                _render_code_block(doc, code_buffer)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Table handling
        if "|" in line and line.strip().startswith("|"):
            table_rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                if re.match(r'^\|[\s\-:|]+\|$', lines[i].strip()):
                    i += 1
                    continue
                cells = [c.strip() for c in lines[i].strip().split("|")[1:-1]]
                table_rows.append(cells)
                i += 1
            if table_rows:
                _render_table(doc, table_rows)
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            p = doc.add_paragraph()
            run = p.add_run("_" * 60)
            run.font.color.rgb = RGBColor(200, 200, 200)
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # Blockquote
        if line.strip().startswith(">"):
            quote_text = line.strip().lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(quote_text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        # Image (add placeholder text)
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            alt_text = img_match.group(1) or "Image"
            p = doc.add_paragraph()
            run = p.add_run(f"[Image: {alt_text}]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            i += 1
            continue

        # List items (numbered)
        num_match = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if num_match:
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, num_match.group(2))
            i += 1
            continue

        # List items (bullet)
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            bullet_text = line.strip()[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, bullet_text)
            i += 1
            continue

        # Italic-only line (image caption)
        stripped = line.strip()
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            caption = stripped.strip("*")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, line.strip())
        i += 1

    # Footer with generation timestamp
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from datetime import datetime
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M')
    run = p.add_run(f"Generated by Consensus Pipeline | {now_str}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def _add_formatted_text(paragraph, text: str):
    """Add text with bold/italic/code formatting."""
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `code`
            run = paragraph.add_run(match.group(4))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(50, 50, 50)
        last_end = match.end()
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _render_code_block(doc, lines: List[str]):
    """Render code block with monospace font and gray background."""
    for code_line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(code_line if code_line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(50, 50, 50)
        # Gray background shading
        shading_elm = p._element.get_or_add_pPr()
        shd = shading_elm.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "F5F5F5",
        })
        shading_elm.append(shd)


def _render_table(doc, rows: List[List[str]]):
    """Render markdown table as Word table."""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row[:num_cols]):
            cell = table.cell(i, j)
            cell.text = cell_text
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()  # spacing after table
