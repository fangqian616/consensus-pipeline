"""
AI Consensus Pipeline v4.0 - Multi-department AI debate content creation framework
v4.0: Requirement research tab (Phase 0-4) + Programming/Tutorial departments + Academic research integration
Supports step-by-step mode: pause after each debate round, director reviews and gives correction instructions
v3.0: Consensus Pipeline restructure — Smart grouping tab + preset/user config management + AI Router auto-grouping
v2.4: Market mode + candidate campaign + vote election + patch correction
v2.2: Screenwriter dept adds narrative architect + DP switch motivation + storyboard tabularization + browser notifications
v2.1: Spatial positioning restructure: from coordinate anchor to object positioning + frame position dual-layer structure
"""
import streamlit as st
import streamlit.components.v1 as st_components
import json
import os
import time
import uuid
import tempfile
from datetime import datetime


from requirement.fact_checker import FactChecker

from debate_engine import (
    DEPARTMENTS, P2_CROSS_DEBATES, P5_CROSS_DEBATES, CROSS_DEBATES,
    STRUCTURED_TEMPLATES, PROOFREAD_DEPARTMENTS, DEPT_ORDER,
    ACADEMIC_PROOFREAD_DEPARTMENTS,
    MODEL_PROFILES, ARCHITECTURE_MODES,
    MARKET_CONFIG,
    apply_config, get_current_config, get_current_config_name,
    run_department_debate, run_cross_debate, run_summary, run_academic_summary,
    run_academic_search, build_papers_summary,
    _extract_search_strategy, _regenerate_strategy_json,
    LIT_STRATEGY_INSTRUCTION_ZH, LIT_STRATEGY_INSTRUCTION_EN,
    run_spatial_review, run_proofreading, run_academic_proofreading, run_academic_auto_revision, run_spatial_diagram,
    run_auto_revision, run_director_revision, run_output_edit,
    run_smart_reroll,
    run_department_round, run_department_consensus,
    generate_carry_forward, generate_asset_checklist,
    run_single_agent, run_expert_pool_debate,
    generate_candidates, generate_questions, vote_on_questions, patch_winner,
    clean_spatial_coordinates,
)
from config_manager import (
    list_presets, load_preset, list_profiles, save_profile, load_profile,
    delete_profile, export_config, import_config,
    get_last_used, set_last_used, validate_config, merge_skill_injection,
)
from router import analyze_and_configure, analyze_revision_impact

# v4.0 Requirement research module
from requirement import (

    RequirementDocument, RequirementInterviewer,
    StructuredRequirement, RequirementStructurer,
    DiscussionGroup, DiscussionResult,
    ConfigRecommender,
    FactChecker, FactCheckReport, FactCheckResult,
)

# ============ ASCII-safe filename helper ============
def _safe_filename(name: str, fallback: str = "download") -> str:
    """Sanitize filename to ASCII-safe chars for HTTP Content-Disposition header.
    Streamlit's download_button sets Content-Disposition which uses latin-1 encoding;
    non-ASCII chars (Chinese etc.) cause UnicodeEncodeError.
    """
    import re as _re
    safe = _re.sub(r'[^\x20-\x7e]', '_', name)  # replace non-ASCII with _
    safe = _re.sub(r'_+', '_', safe).strip('_')   # collapse multiple _
    return safe if safe else fallback

# ============ Pipeline Mode Detection ============

def _detect_pipeline_mode(cfg=None):
    """Detect pipeline mode: 'academic' or 'animation'.
    Priority: session_state.pipeline_mode > department key detection.
    """
    explicit = st.session_state.get("pipeline_mode")
    if explicit in ("academic", "animation"):
        return explicit
    # Fallback: detect from department keys
    if cfg is None:
        cfg = (st.session_state.get("workgroup_config") or get_current_config() or {})
    dept_keys = list(cfg.get("departments", {}).keys()) if isinstance(cfg, dict) else []
    academic_keys = {"literature_search", "methodology_review", "report_integration", 
                     "programming", "tutorial", "metadata_inspector", "citation_network",
                     "data_validation", "counter_evidence", "topic_clustering", "visualization"}
    if any(k in dept_keys for k in academic_keys):
        return "academic"
    return "animation"


# ============ Browser Notifications ============

# Notification sound: base64-encoded short chime (~0.3s ding-dong)
_NOTIFICATION_AUDIO_B64 = """
data:audio/wav;base64,UklGRnoGAABXQVZFZm10IBAAAAABAAEAQB8AAEAfAAABAAgAZGF0YQoGAACBhYqFbF1fdH2LkI2Hg3x1b2Vpe4eQlZONhX54cGhgYGl1gIuSlZKMhX54c2xkYWdwe4SOk5WRjIV+eHNsZGFncHuEjpOVkYyFfnhzbGRhZ3B7hI6TlZGMhX54c2xkYWdwe4SOk5WRjIV+eHNsZGFncHs=
"""

def browser_notify(title: str, message: str):
    """
    Send browser desktop notification + play chime.
    Works even when Streamlit page is in a background tab.
    """
    # Desktop notification
    notify_js = f"""
    <script>
    (function() {{
        if (!("Notification" in window)) return;
        if (Notification.permission === "granted") {{
            new Notification("{title}", {{ body: "{message}", icon: "🎬" }});
        }} else if (Notification.permission !== "denied") {{
            Notification.requestPermission().then(function(perm) {{
                if (perm === "granted") {{
                    new Notification("{title}", {{ body: "{message}", icon: "🎬" }});
                }}
            }});
        }}
    }})();
    </script>
    """
    st_components.html(notify_js, height=0)
    
    # Notification sound
    audio_html = f"""
    <audio autoplay style="display:none">
    <source src="{_NOTIFICATION_AUDIO_B64}" type="audio/wav">
    </audio>
    """
    st_components.html(audio_html, height=0)


# ============ Disk Persistence ============
# Auto-save to session_state (primary) + per-user file backup (recovery safety net).
# File backup survives session_state loss within same container lifecycle
# (e.g. rerun error, script crash). Does NOT survive container restart (OOM, git push).

_AUTOSAVE_PREFIX = "_autosave_"
_autosave_dir = os.path.join(tempfile.gettempdir(), "consensus_autosave")

def _autosave_file_path(key: str) -> str:
    """Per-user file path for autosave backup. UUID in session_state makes filename unguessable."""
    if "_session_uuid" not in st.session_state:
        st.session_state._session_uuid = uuid.uuid4().hex
    return os.path.join(_autosave_dir, f"{st.session_state._session_uuid}_{key}.json")

def autosave_result(key: str, data: dict):
    """
    Persist results to session_state (primary) + file backup (recovery).
    key: e.g. "normal_result", "market_step1", "market_result"
    data: dict data to save
    """
    try:
        st.session_state[_AUTOSAVE_PREFIX + key] = data
    except Exception as e:
        import sys
        print(f"[autosave] Save {key} to session_state failed: {e}", file=sys.stderr)
    # File backup (best-effort, non-blocking)
    try:
        os.makedirs(_autosave_dir, exist_ok=True)
        with open(_autosave_file_path(key), "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
    except Exception:
        pass  # File backup is optional; don't break main flow

def autosave_load(key: str) -> dict | None:
    """
    Load persisted results from session_state.
    Falls back to file backup if session_state is empty (recovery scenario).
    Returns dict or None
    """
    result = st.session_state.get(_AUTOSAVE_PREFIX + key)
    if result is not None:
        return result
    # Fallback: try file backup (session_state was lost but file survived)
    try:
        path = _autosave_file_path(key)
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return None

def autosave_has(key: str) -> bool:
    """Check if persisted result exists (in session_state or file backup)"""
    if (_AUTOSAVE_PREFIX + key) in st.session_state:
        return True
    try:
        return os.path.exists(_autosave_file_path(key))
    except Exception:
        return False

def autosave_list() -> list[str]:
    """List all persisted result keys for current user"""
    return [
        k[len(_AUTOSAVE_PREFIX):]
        for k in st.session_state
        if k.startswith(_AUTOSAVE_PREFIX)
    ]

def autosave_clear(key: str = None):
    """Clear persisted results. key=None clears all"""
    if key:
        st.session_state.pop(_AUTOSAVE_PREFIX + key, None)
        try:
            path = _autosave_file_path(key)
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass
    else:
        for k in list(st.session_state):
            if k.startswith(_AUTOSAVE_PREFIX):
                del st.session_state[k]
        # Clean file backups for this session
        try:
            uuid_str = st.session_state.get("_session_uuid", "")
            if uuid_str and os.path.isdir(_autosave_dir):
                for f in os.listdir(_autosave_dir):
                    if f.startswith(uuid_str):
                        os.remove(os.path.join(_autosave_dir, f))
        except Exception:
            pass


def notify_stage_complete(stage_name: str, detail: str = ""):
    """Shortcut for stage completion notification"""
    is_zh = st.session_state.get("lang", "zh") == "zh"
    title = "🎬 辩论系统" if is_zh else "🎬 Debate System"
    msg = f"✅ {stage_name} complete!" if is_zh else f"✅ {stage_name} complete!"
    if detail:
        msg += f" {detail}"
    browser_notify(title, msg)


# ============ i18n ============
LANG = {
    "zh": {
        "title": "🧠 AI Consensus Pipeline",
        "subtitle": "通用AI多部门辩论内容创作框架 — 智能配组 → 部门辩论 → 交叉审核 → 校对 → 产出",
        "tab_input": "📝 输入",
        "tab_debate": "🗣️ 部门辩论",
        "tab_cross": "⚔️ 交叉辩论",
        "tab_output": "🎬 最终产出",
        "tab_proofread": "🔍 校对",
        "script_label": "基础剧本",
        "script_hint": "输入你的场景、角色、情节骨架",
        "positive_prompt_label": "场景正向提示词",
        "positive_prompt_hint": "视觉风格、渲染质感、布光风格、氛围",
        "negative_prompt_label": "负面提示词",
        "negative_prompt_hint": "要排除的视觉元素",
        "character_refs_label": "角色参考",
        "character_refs_hint": "角色三视图/表情设计描述或参考图提示词",
        "start_debate": "🚀 一键全辩（自动跑完全部）",
        "start_step": "🎬 逐步辩论（每轮暂停审阅）",
        "debate_progress": "辩论进度",
        "consensus": "最终共识",
        "cross_debate_title": "交叉辩论",
        "p2_cross_title": "P2 空间板块交叉审核",
        "p5_cross_title": "P5 交叉辩论",
        "vs": " vs ",
        "output_storyboard": "📐 静态关键帧分镜表（喂绘图工具）",
        "output_video": "🎥 逐镜视频提示词（喂视频生成工具）",
        "copy_prompt": "📋 复制",
        "export_json": "📥 导出完整结果JSON",
        "export_storyboard": "📥 导出分镜表",
        "export_video": "📥 导出视频提示词",
        "debate_rounds": "辩论轮次",
        "debate_rounds_hint": "每个部门的辩论轮次（1-10），轮次越多讨论越深入",
        "api_url": "API地址",
        "api_key": "API密钥",
        "model_name": "模型名称",
        "language": "语言 / Language",
        "need_api": "请在侧边栏配置API信息",
        "need_script": "请先在「输入」页填写基础剧本",
        "debate_not_started": "辩论尚未开始，请先在「输入」页填写剧本并点击开始",
        "cross_not_started": "交叉辩论需要先完成部门辩论",
        "output_not_ready": "最终产出需要先完成全部辩论",
        "proofread_not_ready": "校对需要先完成最终产出",
        "rerun_dept": "🔄 重新辩论此部门",
        "dept_result": "辩论结果",
        "extra_instructions": "额外指令",
        "extra_instructions_hint": "对当前部门的额外创作指令",
        "segment_note": "💡 如剧本超过15秒，产物将自动拆为多段",
        "director_review": "🎬 导演审阅",
        "approve": "✅ 通过",
        "reject_redebate": "❌ 打回重辩",
        "reject_hint": "请输入修改意见后打回",
        "rejection_reason": "修改意见",
        "proofread_run": "🔍 执行校对",
        "proofread_running": "🔍 正在校对...",
        "proofread_result": "校对结果",
        "proofread_overall": "总体评价",
        # Step-by-step mode
        "step_mode": "逐步辩论模式",
        "step_mode_hint": "开启后每轮辩论暂停，导演可审阅并给出纠正指令再继续",
        "step_round_title": "第{round}轮辩论完成",
        "step_correction_label": "导演纠正指令（可选）",
        "step_correction_hint": "如果本轮辩论跑偏，在此输入纠正指令，将在下一轮生效",
        "step_continue": "▶️ 继续下一轮",
        "step_finish_dept": "✅ 结束辩论，生成共识",
        "step_dept_done": "{dept} 辩论完成",
        "step_next_dept": "▶️ 开始{dept}辩论",
        "step_all_done": "✅ 全部部门辩论完成",
        "step_run_cross": "⚔️ 开始交叉辩论",
        "step_run_summary": "🎬 生成最终产出",
        "step_run_proofread": "🔍 执行校对",
        "step_waiting": "等待导演操作",
        "step_current_status": "当前状态",
        # Carry-forward document
        "carry_forward_label": "📋 承上文档（前段决策摘要）",
        "carry_forward_hint": "粘贴上一段辩论生成的承上文档，保证多段剧本之间的连续性",
        "carry_forward_generate": "🔄 生成承上文档",
        "carry_forward_generated": "✅ 承上文档已生成",
        "carry_forward_copy": "📋 复制承上文档",
        "carry_forward_export": "📥 导出承上文档",
        "carry_forward_note": "💡 辩论完成后点击「生成承上文档」，将产物粘贴到下一段的承上文档输入框中",
        # Spatial diagram
        "spatial_diagram": "🗺️ 空间示意图",
        "spatial_diagram_generate": "🗺️ 生成空间示意图提示词",
        "spatial_diagram_generating": "🗺️ 正在生成空间示意图提示词...",
        "spatial_diagram_result": "空间示意图提示词",
        "spatial_diagram_export": "📥 导出空间示意图提示词",
        "spatial_diagram_hint": "根据空间板块共识生成可视化布局图提示词，可喂给AI绘图工具生成场景俯视图",
        "spatial_diagram_need_spatial": "空间示意图需要先完成空间板块辩论",
        "auto_revision": "🔧 自动修正",
        "auto_revision_run": "🔧 根据校对反馈自动修正",
        "auto_revision_running": "🔧 正在根据校对反馈修正产出...",
        "auto_revision_done": "✅ 自动修正完成",
        "auto_revision_notes": "修正说明",
        "director_revision": "🎬 导演指令修正",
        "director_revision_dept": "选择要重辩的部门",
        "director_revision_note": "修改意见",
        "director_revision_note_hint": "告诉这个部门哪里不对、要怎么改",
        "director_revision_run": "🎬 执行导演指令修正",
        "director_revision_running": "🎬 正在重辩{dept}并重新生成产出...",
        "director_revision_done": "✅ 导演指令修正完成",
        "revision_rounds": "重辩轮次",
        # Output re-roll editing
        "output_edit": "✏️ 产出回炉编辑",
        "output_edit_hint": "选中部门+给修改意见 → 从该部门专业视角定向修改分镜表/视频提示词/空间表（不重跑辩论）",
        "output_edit_dept": "选择编辑部门",
        "output_edit_note": "修改意见",
        "output_edit_note_hint": "告诉这个部门哪里要改、要加什么内容",
        "output_edit_spatial": "同时修改空间板块定义（物品清单/角色定位/动线）",
        "output_edit_run": "✏️ 执行回炉编辑",
        "output_edit_running": "✏️ 正在让{dept}编辑产出...",
        "output_edit_done": "✅ 回炉编辑完成",
        "output_edit_notes": "修改说明",
        "output_edit_apply": "✅ 应用编辑结果",
        "output_edit_applied": "✅ 编辑结果已应用到产出",
        "output_edit_spatial_updated": "（空间板块定义已同步更新）",
        # Model and architecture
        "model_profile": "🤖 模型配置",
        "model_profile_hint": "选择预置模型或自定义",
        "custom_api_url": "自定义API地址",
        "custom_api_key": "自定义API密钥",
        "custom_model_name": "自定义模型名称",
        "architecture_mode": "🏗️ 架构模式",
        "mode_pipeline": "共识管线 (Pipeline of Consensus)",
        "mode_single_agent": "单Agent基线 (Baseline)",
        "mode_expert_pool": "专家池 (Expert Pool)",
        # Comparison panel
        "tab_compare": "📊 对比",
        "compare_title": "📊 运行对比面板",
        "compare_desc": "同一剧本，不同模型/架构的运行结果对比。展示Harness Engineering的核心论点：换架构 > 换模型。",
        "compare_no_runs": "还没有运行记录。跑一次辩论后，结果会自动保存到这里。",
        "compare_save_run": "💾 保存当前运行结果",
        "compare_saved": "✅ 运行结果已保存",
        "compare_run_label": "运行 #{idx}",
        "compare_model": "模型",
        "compare_mode": "架构",
        "compare_tokens": "总Token",
        "compare_time": "耗时",
        "compare_api_calls": "API调用次数",
        "compare_storyboard_preview": "分镜表预览",
        "compare_video_preview": "视频提示词预览",
        "compare_delete": "🗑️ 删除此运行",
        "compare_clear": "🗑️ 清空所有运行",
        "compare_table": "对比总览",
        "mode_pipeline_desc": "8部门 × 3-4辩手 × N轮 → 共识管线",
        "mode_single_desc": "1次LLM调用 → 无结构基线",
        "mode_expert_desc": "场景分析 → 2辩手/部门 → 精简管线",
        # Market mode
        "tab_market": "🏪 市场",
        "market_title": "🏪 市场模式",
        "market_subtitle": "3个候选竞选 → 100题投票选举 → 补丁修正 → 最优产出",
        "market_num_candidates": "候选数量",
        "market_num_candidates_hint": "生成几个候选方案（2-5），每个用不同温度",
        "market_questions_per": "每辩手出题数",
        "market_questions_per_hint": "每个AI提出几个评估问题（3-15）",
        "market_start": "🏪 开市！",
        "market_running": "🏪 市场营业中...",
        "market_step1": "📦 Step 1: 生成候选方案",
        "market_step2": "❓ Step 2: 出题",
        "market_step3": "🗳️ Step 3: 投票",
        "market_step4": "🔧 Step 4: 补丁修正",
        "market_generating": "正在生成候选{idx}/{total}（温度{temp}）...",
        "market_questioning": "正在出题：{dept}/{debater}...",
        "market_voting": "正在投票：{voter}...",
        "market_patching": "正在补丁修正...",
        "market_candidates": "候选方案",
        "market_questions_total": "共{total}个问题",
        "market_vote_result": "投票结果",
        "market_winner": "🏆 当选方案：候选{label}（{votes}票）",
        "market_vote_detail": "投票详情",
        "market_contested": "争议问题",
        "market_contested_hint": "投票分歧较大的问题，将用于补丁修正",
        "market_patch_result": "补丁修正结果",
        "market_patch_notes": "修正说明",
        "market_apply_winner": "✅ 应用当选方案",
        "market_apply_patched": "✅ 应用补丁修正版",
        "market_applied": "✅ 方案已应用到最终产出",
        "market_no_result": "市场尚未开市，请先填写剧本并点击「开张！」",
        "market_need_debate": "请先在输入页完成剧本填写和API配置",
        "market_parallel_workers": "并行数量",
        "market_parallel_workers_hint": "同时生成几个候选（1=串行，3=3个一起跑）",
        "market_rounds": "市场辩论轮次",
        "market_rounds_hint": "每个候选内部辩论轮次，与普通模式对齐以保证公平对比",
        "asset_checklist": "📋 资产参照表",
        "asset_checklist_hint": "从各部门共识自动提取制作动画前需要准备的素材清单",
        "asset_checklist_download": "📥 导出资产参照表",
        # v3.0 Smart grouping
        "tab_config": "🧠 智能配组",
        "config_title": "🧠 Consensus Pipeline 智能配组",
        "config_subtitle": "描述你的创作需求，AI自动配置最优辩论工作组",
        "config_input_label": "描述你想制作的内容...",
        "config_input_hint": "例如：制作一个2分钟的科幻短片、写一篇悬疑小说、设计一个RPG游戏的战斗系统...",
        "config_ai_button": "🧠 AI智能配组",
        "config_ai_running": "🧠 AI正在分析你的需求并配置工作组...",
        "config_preset_label": "预设配置",
        "config_profile_label": "我的配置",
        "config_preview_title": "工作组预览",
        "config_dept_enabled": "启用",
        "config_dept_disabled": "禁用",
        "config_debaters": "辩手",
        "config_global_params": "全局参数",
        "config_visual_directive": "视觉指令",
        "config_negative_prompts": "负面提示词",
        "config_debate_rounds": "辩论轮次",
        "config_skill_injection": "Skill注入",
        "config_skill_injection_hint": "输入markdown内容，将自动注入到相关辩手提示词末尾",
        "config_skill_target": "注入目标部门",
        "config_skill_target_all": "所有部门",
        "config_save": "💾 保存配置",
        "config_save_name": "配置名称",
        "config_export": "📋 导出配置",
        "config_import": "📥 导入配置",
        "config_apply": "✅ 确认并开始",
        "config_applied": "✅ 配置已应用！切换到输入页开始创作",
        "config_delete": "🗑️ 删除",
        "config_no_depts": "尚未配置任何部门，请使用AI配组或选择预设",
        "config_clarification": "⚠️ AI需要更多信息",
        "config_error": "❌ 配置出错",
        "config_current": "当前配置",
        # --- Tab restructure keys ---
        "tab_setup": "🔬 需求与配置",
        "tab_setup_req": "📋 需求调研",
        "tab_setup_config": "🧠 智能配组",
        "tab_setup_input": "📝 内容输入",
        "tab_debate_combined": "🗣️ 辩论",
        "tab_dept_debate": "🗣️ 部门辩论",
        "tab_cross_debate_sub": "⚔️ 交叉辩论",
        "tab_output_combined": "🎬 产出",
        "tab_final_output": "🎬 最终产出",
        "tab_proofread_sub": "🔍 校对",
        "tab_tools": "📊 工具",
        "tab_compare_sub": "📊 运行对比",
        "tab_market_sub": "🏪 市场模式",
        "toast_config_confirmed": "✅ 配置已确认！请切换到「辩论」Tab开始",
    },
    "en": {
        "title": "🧠 AI Consensus Pipeline",
        "subtitle": "Universal AI Multi-Dept Debate Framework — Smart Config → Dept Debate → Cross Review → Proofread → Output",
        "tab_input": "📝 Input",
        "tab_debate": "🗣️ Dept. Debate",
        "tab_cross": "⚔️ Cross Debate",
        "tab_output": "🎬 Output",
        "tab_proofread": "🔍 Proofread",
        "script_label": "Base Script",
        "script_hint": "Enter your scene, characters, and plot skeleton",
        "positive_prompt_label": "Scene Positive Prompt",
        "positive_prompt_hint": "Visual style, render quality, lighting style, atmosphere",
        "negative_prompt_label": "Negative Prompt",
        "negative_prompt_hint": "Visual elements to exclude",
        "character_refs_label": "Character References",
        "character_refs_hint": "Character turnaround/expression design descriptions or reference image prompts",
        "start_debate": "🚀 Auto Run All",
        "start_step": "🎬 Step-by-Step (pause each round)",
        "debate_progress": "Debate Progress",
        "consensus": "Final Consensus",
        "cross_debate_title": "Cross-Department Debate",
        "p2_cross_title": "P2 Spatial Review",
        "p5_cross_title": "P5 Cross-Debate",
        "vs": " vs ",
        "output_storyboard": "📐 Static Keyframe Storyboard (for image generation)",
        "output_video": "🎥 Per-Shot Video Prompt (for video generation)",
        "copy_prompt": "📋 Copy",
        "export_json": "📥 Export Full Results JSON",
        "export_storyboard": "📥 Export Storyboard",
        "export_video": "📥 Export Video Prompts",
        "debate_rounds": "Debate Rounds",
        "debate_rounds_hint": "Rounds per department (1-10). More rounds = deeper discussion",
        "api_url": "API URL",
        "api_key": "API Key",
        "model_name": "Model Name",
        "language": "语言 / Language",
        "need_api": "Please configure API in the sidebar",
        "need_script": "Please fill in the base script on the Input tab first",
        "debate_not_started": "Debate not started. Fill in the script and click Start.",
        "cross_not_started": "Complete department debates first",
        "output_not_ready": "Complete all debates first",
        "proofread_not_ready": "Generate final output first",
        "rerun_dept": "🔄 Re-debate this department",
        "dept_result": "Debate Results",
        "extra_instructions": "Extra Instructions",
        "extra_instructions_hint": "Additional creative instructions for current department",
        "segment_note": "💡 If script exceeds 15s, output will be split into segments",
        "director_review": "🎬 Director Review",
        "approve": "✅ Approve",
        "reject_redebate": "❌ Reject & Re-debate",
        "reject_hint": "Enter revision notes before rejecting",
        "rejection_reason": "Revision Notes",
        "proofread_run": "🔍 Run Proofreading",
        "proofread_running": "🔍 Proofreading...",
        "proofread_result": "Proofreading Result",
        "proofread_overall": "Overall Assessment",
        # Step mode
        "step_mode": "Step-by-Step Mode",
        "step_mode_hint": "Pause after each round; director can review and give corrections before continuing",
        "step_round_title": "Round {round} Complete",
        "step_correction_label": "Director Correction (optional)",
        "step_correction_hint": "If this round went off track, enter corrections here for the next round",
        "step_continue": "▶️ Continue to Next Round",
        "step_finish_dept": "✅ End Debate & Generate Consensus",
        "step_dept_done": "{dept} Debate Complete",
        "step_next_dept": "▶️ Start {dept} Debate",
        "step_all_done": "✅ All Department Debates Complete",
        "step_run_cross": "⚔️ Start Cross-Debate",
        "step_run_summary": "🎬 Generate Final Output",
        "step_run_proofread": "🔍 Run Proofreading",
        "step_waiting": "Waiting for director action",
        "step_current_status": "Current Status",
        # Carry forward
        "carry_forward_label": "📋 Carry Forward (Previous Segment Summary)",
        "carry_forward_hint": "Paste the carry-forward document from the previous segment's debate to ensure continuity",
        "carry_forward_generate": "🔄 Generate Carry Forward",
        "carry_forward_generated": "✅ Carry Forward Generated",
        "carry_forward_copy": "📋 Copy Carry Forward",
        "carry_forward_export": "📥 Export Carry Forward",
        "carry_forward_note": "💡 After debate completes, click 'Generate Carry Forward' and paste it into the next segment's input",
        # Spatial diagram
        "spatial_diagram": "🗺️ Spatial Diagram",
        "spatial_diagram_generate": "🗺️ Generate Spatial Diagram Prompt",
        "spatial_diagram_generating": "🗺️ Generating spatial diagram prompt...",
        "spatial_diagram_result": "Spatial Diagram Prompt",
        "spatial_diagram_export": "📥 Export Spatial Diagram Prompt",
        "spatial_diagram_hint": "Generate a visual layout diagram prompt from Spatial Planning consensus, for AI image tools to create scene top-down view",
        "spatial_diagram_need_spatial": "Spatial diagram requires Spatial Planning debate to be completed first",
        "auto_revision": "🔧 Auto Revision",
        "auto_revision_run": "🔧 Auto-fix based on proofreading feedback",
        "auto_revision_running": "🔧 Revising output based on proofreading feedback...",
        "auto_revision_done": "✅ Auto revision complete",
        "auto_revision_notes": "Revision Notes",
        "director_revision": "🎬 Director Revision",
        "director_revision_dept": "Select department to re-debate",
        "director_revision_note": "Revision Notes",
        "director_revision_note_hint": "Tell this department what's wrong and how to fix it",
        "director_revision_run": "🎬 Execute Director Revision",
        "director_revision_running": "🎬 Re-debating {dept} and regenerating output...",
        "director_revision_done": "✅ Director revision complete",
        "revision_rounds": "Re-debate Rounds",
        # Output edit
        "output_edit": "✏️ Output Revision",
        "output_edit_hint": "Select department + edit instructions → targeted edit of storyboard/video prompt/spatial layout from that department's perspective (no re-debate)",
        "output_edit_dept": "Select editing department",
        "output_edit_note": "Edit Instructions",
        "output_edit_note_hint": "Tell this department what to change or add",
        "output_edit_spatial": "Also edit spatial layout (object inventory/character positioning/paths)",
        "output_edit_run": "✏️ Execute Revision",
        "output_edit_running": "✏️ {dept} is editing the output...",
        "output_edit_done": "✅ Output revision complete",
        "output_edit_notes": "Edit Notes",
        "output_edit_apply": "✅ Apply Edit Results",
        "output_edit_applied": "✅ Edit results applied to output",
        "output_edit_spatial_updated": "(Spatial layout updated)",
        # Model & Architecture
        "model_profile": "🤖 Model Config",
        "model_profile_hint": "Select a preset model or customize",
        "custom_api_url": "Custom API URL",
        "custom_api_key": "Custom API Key",
        "custom_model_name": "Custom Model Name",
        "architecture_mode": "🏗️ Architecture Mode",
        "mode_pipeline": "Pipeline of Consensus",
        "mode_single_agent": "Single Agent Baseline",
        "mode_expert_pool": "Expert Pool",
        # Compare panel
        "tab_compare": "📊 Compare",
        "compare_title": "📊 Run Comparison Panel",
        "compare_desc": "Same script, different model/architecture runs. Demonstrates the core Harness Engineering thesis: architecture > model.",
        "compare_no_runs": "No runs recorded yet. Run a debate and results will be saved here automatically.",
        "compare_save_run": "💾 Save Current Run",
        "compare_saved": "✅ Run results saved",
        "compare_run_label": "Run #{idx}",
        "compare_model": "Model",
        "compare_mode": "Architecture",
        "compare_tokens": "Total Tokens",
        "compare_time": "Duration",
        "compare_api_calls": "API Calls",
        "compare_storyboard_preview": "Storyboard Preview",
        "compare_video_preview": "Video Prompt Preview",
        "compare_delete": "🗑️ Delete Run",
        "compare_clear": "🗑️ Clear All Runs",
        "compare_table": "Comparison Overview",
        "mode_pipeline_desc": "8 depts × 3-4 debaters × N rounds → Consensus Pipeline",
        "mode_single_desc": "1 LLM call → Unstructured baseline",
        "mode_expert_desc": "Scene analysis → 2 debaters/dept → Streamlined Pipeline",
        # Market mode
        "tab_market": "🏪 Market",
        "market_title": "🏪 Market Mode",
        "market_subtitle": "3 candidates run → 100-question vote → patch revision → best output",
        "market_num_candidates": "Number of Candidates",
        "market_num_candidates_hint": "How many candidates to generate (2-5), each with different temperature",
        "market_questions_per": "Questions per Debater",
        "market_questions_per_hint": "How many evaluation questions each AI proposes (3-15)",
        "market_start": "🏪 Market Open!",
        "market_running": "🏪 Market is running...",
        "market_step1": "📦 Step 1: Generate Candidates",
        "market_step2": "❓ Step 2: Generate Questions",
        "market_step3": "🗳️ Step 3: Vote",
        "market_step4": "🔧 Step 4: Patch Revision",
        "market_generating": "Generating candidate {idx}/{total} (temp {temp})...",
        "market_questioning": "Questioning: {dept}/{debater}...",
        "market_voting": "Voting: {voter}...",
        "market_patching": "Patching winner...",
        "market_candidates": "Candidates",
        "market_questions_total": "{total} questions total",
        "market_vote_result": "Vote Results",
        "market_winner": "🏆 Winner: Candidate {label} ({votes} votes)",
        "market_vote_detail": "Vote Details",
        "market_contested": "Contested Issues",
        "market_contested_hint": "Questions with split votes, used for patch revision",
        "market_patch_result": "Patch Revision Result",
        "market_patch_notes": "Patch Notes",
        "market_apply_winner": "✅ Apply Winner",
        "market_apply_patched": "✅ Apply Patched Version",
        "market_applied": "✅ Applied to final output",
        "market_no_result": "Market hasn't opened yet. Fill in the script and click 'Open!'",
        "market_need_debate": "Please fill in the script and API config first",
        "market_parallel_workers": "Parallel Workers",
        "market_parallel_workers_hint": "Generate candidates simultaneously (1=serial, 3=3 at once)",
        "market_rounds": "Market Debate Rounds",
        "market_rounds_hint": "Debate rounds per candidate, aligned with normal mode for fair comparison",
        "asset_checklist": "📋 Asset Checklist",
        "asset_checklist_hint": "Auto-extract asset preparation list from department consensuses before animation production",
        "asset_checklist_download": "📥 Export Asset Checklist",
        # v3.0 Smart Config
        "tab_config": "🧠 Smart Config",
        "config_title": "🧠 Consensus Pipeline Smart Config",
        "config_subtitle": "Describe your creative needs, AI auto-configures the optimal debate workgroup",
        "config_input_label": "Describe what you want to create...",
        "config_input_hint": "e.g., produce a 2-min sci-fi short, write a mystery novel, design an RPG combat system...",
        "config_ai_button": "🧠 AI Smart Config",
        "config_ai_running": "🧠 AI is analyzing your needs and configuring workgroup...",
        "config_preset_label": "Presets",
        "config_profile_label": "My Configs",
        "config_preview_title": "Workgroup Preview",
        "config_dept_enabled": "Enabled",
        "config_dept_disabled": "Disabled",
        "config_debaters": "Debaters",
        "config_global_params": "Global Parameters",
        "config_visual_directive": "Visual Directive",
        "config_negative_prompts": "Negative Prompts",
        "config_debate_rounds": "Debate Rounds",
        "config_skill_injection": "Skill Injection",
        "config_skill_injection_hint": "Enter markdown content, will be auto-injected to relevant debater prompts",
        "config_skill_target": "Target Departments",
        "config_skill_target_all": "All Departments",
        "config_save": "💾 Save Config",
        "config_save_name": "Config Name",
        "config_export": "📋 Export Config",
        "config_import": "📥 Import Config",
        "config_apply": "✅ Confirm & Start",
        "config_applied": "✅ Config applied! Switch to Input tab to start creating",
        "config_delete": "🗑️ Delete",
        "config_no_depts": "No departments configured yet, use AI config or select a preset",
        "config_clarification": "⚠️ AI needs more information",
        "config_error": "❌ Configuration error",
        "config_current": "Current Config",
        # --- Tab restructure keys ---
        "tab_setup": "🔬 Setup",
        "tab_setup_req": "📋 Requirement",
        "tab_setup_config": "🧠 Smart Config",
        "tab_setup_input": "📝 Input",
        "tab_debate_combined": "🗣️ Debate",
        "tab_dept_debate": "🗣️ Dept. Debate",
        "tab_cross_debate_sub": "⚔️ Cross Debate",
        "tab_output_combined": "🎬 Output",
        "tab_final_output": "🎬 Final Output",
        "tab_proofread_sub": "🔍 Proofread",
        "tab_tools": "📊 Tools",
        "tab_compare_sub": "📊 Compare",
        "tab_market_sub": "🏪 Market",
        "toast_config_confirmed": "✅ Config confirmed! Switch to Debate tab",
    }
}

def t(key: str, **kwargs) -> str:
    lang = st.session_state.get("lang", "zh")
    text = LANG.get(lang, LANG["zh"]).get(key, key)
    if kwargs:
        text = text.format(**kwargs)
    return text



def init_state():
    defaults = {
        "lang": "zh",
        "script": "",
        "positive_prompt": "",
        "negative_prompt": "",
        "character_refs": "",
        "api_url": "https://api.deepseek.com/v1/chat/completions",
        "api_key": "",
        "model_name": "deepseek-v4-flash",
        "debate_rounds": 3,
        "extra_instructions": "",
        "step_mode": True,
        # Debate result
        "dept_results": {},
        "spatial_review_result": None,
        "cross_results": [],
        "final_output": {},
        "proofread_result": None,
        "debate_completed": False,
        # Step-by-step mode state
        "step_phase": "idle",  # idle / round_done / dept_done / all_dept_done / spatial_review / cross_done / completed
        "step_dept_index": 0,
        "step_round": 1,
        "step_all_arguments": [],
        "step_debate_log": [],
        "step_corrections": [],
        # Carry-forward document
        "carry_forward": "",
        "carry_forward_result": "",
        "auto_revision_result": None,
        "output_edit_result": None,
        "smart_reroll_impact": None,
        "smart_reroll_result": None,
        # v2.3 Harness Engineering
        "model_profile": "deepseek-v4-flash",
        "custom_api_url": "",
        "custom_api_key": "",
        "custom_model_name": "",
        "architecture_mode": "pipeline_of_consensus",
        "run_history": [],  # Comparison panel的运行历史
        "current_stats": None,  # 当前运行的token统计
        "current_start_time": None,  # 当前运行的开始时间
        "current_end_time": None,  # 当前运行的结束时间
        "_debate_running": False,  # 辩论进行中（同步执行，用于禁用tab切换）
        "_debate_fingerprint": None,  # v0.10: 配置指纹，用于rerun隔离
        "_debate_phase": "departments",  # v0.10: 当前辩论阶段 departments/cross/summary/done
        # v2.4 Market
        "market_result": None,  # Market mode结果
        "market_num_candidates": 3,
        "market_questions_per": 7,
        "market_parallel_workers": 3,
        "market_rounds": 3,
        # v3.0 Smart grouping
        "workgroup_config": None,  # 当前工作组配置（PresetConfig dict）
        "pipeline_mode": None,  # "academic" or "animation", set by Phase4 confirm
        "workgroup_name": "动画辩论",  # 当前配置名称
        "_lang_selected": False,  # Language welcome page flag
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v
    
    # v3.0: Auto-load last used config on startup
    if not st.session_state.get("_config_auto_loaded", False):
        st.session_state._config_auto_loaded = True
        last_used = get_last_used()
        if last_used:
            try:
                # Prefer loading from user_profiles, then presets
                profiles = list_profiles()
                if last_used in profiles:
                    cfg = load_profile(last_used)
                else:
                    cfg = load_preset(last_used)
                apply_config(cfg)
                st.session_state.workgroup_config = cfg
                st.session_state.workgroup_name = last_used
            except Exception as e:
                import streamlit as _st
                _st.toast(f"⚠️ 配置加载失败: {e}", icon="⚠️")



def render_sidebar():
    with st.sidebar:
        is_zh = st.session_state.lang == "zh"
        # Language switch
        # Language switch (compact)
        _lang_opts = {"zh": "🇨🇳 中文", "en": "🇬🇧 English"}
        _cur_lang = st.session_state.lang
        _sel_lang = st.selectbox(
            "Language", options=["zh", "en"],
            format_func=lambda x: _lang_opts[x],
            index=["zh", "en"].index(_cur_lang),
            key="_sidebar_lang_select", label_visibility="collapsed"
        )
        if _sel_lang != _cur_lang:
            st.session_state.lang = _sel_lang
            st.rerun()

        st.divider()
        
        # API configuration
        st.subheader("🔑 API")
        
        # Model configuration
        profile_options = {}
        for pk, pv in MODEL_PROFILES.items():
            profile_options[pk] = pv["zh_name"] if st.session_state.lang == "zh" else pv["en_name"]
        
        selected_profile = st.selectbox(
            t("model_profile"),
            options=list(profile_options.keys()),
            format_func=lambda x: profile_options[x],
            key="model_profile_select",
        )
        st.session_state.model_profile = selected_profile
        
        if selected_profile == "custom":
            st.session_state.custom_api_url = st.text_input(
                t("custom_api_url"),
                value=st.session_state.custom_api_url,
            )
            _custom_key_input = st.text_input(
                t("custom_api_key"),
                value=st.session_state.custom_api_key,
                type="password",
            )
            if _custom_key_input:
                st.session_state.custom_api_key = _custom_key_input
            st.session_state.custom_model_name = st.text_input(
                t("custom_model_name"),
                value=st.session_state.custom_model_name,
            )
            st.session_state.api_url = st.session_state.custom_api_url
            st.session_state.api_key = st.session_state.custom_api_key
            st.session_state.model_name = st.session_state.custom_model_name
        else:
            profile = MODEL_PROFILES[selected_profile]
            st.session_state.api_url = profile["api_url"]
            st.session_state.model_name = profile["model"]
            # API Key still manually entered (password field returns empty on rerun, only update when user actually types)
            _api_key_input = st.text_input(
                t("api_key"),
                value=st.session_state.api_key,
                type="password",
            )
            if _api_key_input:
                st.session_state.api_key = _api_key_input
        
        st.divider()
        
        # Debate parameters
        st.subheader("⚙️ " + ("辩论参数" if st.session_state.lang == "zh" else "Parameters"))
        st.session_state.debate_rounds = st.slider(
            t("debate_rounds"),
            min_value=1, max_value=10, value=st.session_state.debate_rounds,
            help=t("debate_rounds_hint"),
        )
        st.session_state.extra_instructions = st.text_area(
            t("extra_instructions"),
            value=st.session_state.extra_instructions,
            height=80,
            help=t("extra_instructions_hint"),
        )
        
        st.divider()
        
        # Architecture mode
        st.subheader(t("architecture_mode"))
        mode_options = {}
        for mk, mv in ARCHITECTURE_MODES.items():
            mode_options[mk] = mv["zh_name"] if st.session_state.lang == "zh" else mv["en_name"]
        
        selected_mode = st.selectbox(
            t("architecture_mode"),
            options=list(mode_options.keys()),
            format_func=lambda x: mode_options[x],
            key="arch_mode_select",
        )
        st.session_state.architecture_mode = selected_mode
        
        # Mode description
        mode_info = ARCHITECTURE_MODES[selected_mode]
        mode_desc = mode_info["zh_desc"] if st.session_state.lang == "zh" else mode_info["en_desc"]
        st.caption(mode_desc)

        st.divider()
        
        # Market parameters
        st.subheader("🏪 " + ("市场参数" if st.session_state.lang == "zh" else "Market Params"))
        st.session_state.market_num_candidates = st.slider(
            t("market_num_candidates"),
            min_value=2, max_value=5, value=st.session_state.market_num_candidates,
            help=t("market_num_candidates_hint"),
        )
        st.session_state.market_questions_per = st.slider(
            t("market_questions_per"),
            min_value=3, max_value=15, value=st.session_state.market_questions_per,
            help=t("market_questions_per_hint"),
        )
        st.session_state.market_parallel_workers = st.slider(
            t("market_parallel_workers"),
            min_value=1, max_value=5, value=st.session_state.market_parallel_workers,
            help=t("market_parallel_workers_hint"),
        )
        st.session_state.market_rounds = st.slider(
            t("market_rounds"),
            min_value=1, max_value=10, value=st.session_state.get("market_rounds", 3),
            help=t("market_rounds_hint"),
        )

        st.divider()
        
        # Step-by-step mode toggle
        st.session_state.step_mode = st.toggle(
            t("step_mode"),
            value=st.session_state.step_mode,
            help=t("step_mode_hint"),
        )

        st.divider()
        
        # Advanced settings
        with st.expander("🔧 " + ("高级设置" if is_zh else "Advanced"), expanded=False):
            _es_key = st.text_input(
                "easyScholar API Key" + ("（可选，国内用户增强期刊排名）" if is_zh else " (optional, CN users only)"),
                value=st.session_state.get("easyscholar_key", ""),
                type="password",
                help="https://www.easyscholar.cc/ → 用户中心 → 开放API" if is_zh else "Optional. Only useful for Chinese journal rankings.",
            )
            if _es_key:
                st.session_state.easyscholar_key = _es_key
                os.environ["EASYSCHOLAR_SECRET_KEY"] = _es_key
            
            st.markdown("---")
            
            # .env file generator (download, not write to shared disk)
            st.markdown("📄 " + ("CLI配置文件生成" if is_zh else "CLI Config (.env)"))
            if st.session_state.get("api_key"):
                env_content = f"DEEPSEEK_API_KEY={st.session_state.api_key}\n"
                if st.session_state.get("easyscholar_key"):
                    env_content += f"EASYSCHOLAR_SECRET_KEY={st.session_state.easyscholar_key}\n"
                st.download_button(
                    label="⬇️ " + ("下载.env文件" if is_zh else "Download .env file"),
                    data=env_content,
                    file_name=".env",
                    mime="text/plain",
                    key="download_env_btn",
                    use_container_width=True,
                )
            else:
                st.caption(("请先在上方填写API Key" if is_zh else "Fill API Key above first"))
        
        st.divider()
        
        # Archive management (per-user session state)
        saved_keys = autosave_list()
        if saved_keys:
            st.subheader("💾 " + ("存档管理" if st.session_state.lang == "zh" else "Autosave"))
            for key in saved_keys:
                st.caption(f"📄 {key}")
            if st.button("🗑️ " + ("清除全部存档" if st.session_state.lang == "zh" else "Clear all saves"), 
                         use_container_width=True):
                autosave_clear()
                st.rerun()
        
        st.divider()
        
        # v3.0 My configs
        st.subheader(t("config_profile_label"))
        _all_profiles = list_profiles()
        _all_presets = list_presets()
        _config_options = []
        for p in _all_presets:
            _config_options.append(f"📦 {p}")
        for p in _all_profiles:
            _config_options.append(f"💾 {p}")
        if _config_options:
            _current_name = get_current_config_name()
            _sel_idx = 0
            for i, opt in enumerate(_config_options):
                clean = opt.replace("📦 ", "").replace("💾 ", "")
                if clean == _current_name:
                    _sel_idx = i
                    break
            _selected = st.selectbox(
                t("config_current"),
                options=_config_options,
                index=_sel_idx,
            )
            if st.button("✅ " + ("应用" if st.session_state.lang == "zh" else "Apply"), use_container_width=True):
                _clean_name = _selected.replace("📦 ", "").replace("💾 ", "")
                try:
                    if _selected.startswith("📦"):
                        _cfg = load_preset(_clean_name)
                    else:
                        _cfg = load_profile(_clean_name)
                    if _cfg is not None:
                        apply_config(_cfg)
                        st.session_state.workgroup_config = _cfg
                        st.session_state.pipeline_mode = None  # reset, will be detected on next use
                    st.session_state.workgroup_name = _clean_name
                    set_last_used(_clean_name)
                    st.rerun()
                except Exception as e:
                    st.error(str(e))


def build_dept_input(dept_key: str) -> str:
    """Build department inputs (P1→P2→P3→P4 flow)"""
    is_zh = st.session_state.lang == "zh"
    script = st.session_state.script
    positive = st.session_state.positive_prompt
    chars = st.session_state.character_refs
    dept_results = st.session_state.get("dept_results", {})
    
    # P1: Screenwriter department
    if dept_key == "screenwriter":
        if is_zh:
            return f"剧本：\n{script}\n\n场景风格：\n{positive}\n\n角色参考：\n{chars}"
        else:
            return f"Script:\n{script}\n\nScene Style:\n{positive}\n\nCharacter References:\n{chars}"
    
    # P2: Spatial planning (based on screenwriter output)
    elif dept_key == "spatial":
        sw = dept_results.get("screenwriter", {}).get("consensus", script)
        if is_zh:
            return f"编剧部细节填充版剧本：\n{sw}\n\n原始场景风格：\n{positive}"
        else:
            return f"Screenwriter Detailed Script:\n{sw}\n\nOriginal Scene Style:\n{positive}"
    
    # P3: Visual four departments (based on screenwriter + spatial output)
    elif dept_key == "storyboard":
        sw = dept_results.get("screenwriter", {}).get("consensus", script)
        sp = dept_results.get("spatial", {}).get("consensus", "")
        if is_zh:
            return f"编剧部细节填充版剧本：\n{sw}\n\n空间板块布局：\n{sp}\n\n场景风格：\n{positive}"
        else:
            return f"Screenwriter Detailed Script:\n{sw}\n\nSpatial Layout:\n{sp}\n\nScene Style:\n{positive}"
    
    elif dept_key in ("dp", "lighting", "vfx"):
        parts = []
        sw_label = "编剧部：" if is_zh else "Screenwriter:"
        parts.append(f"{sw_label}\n{dept_results.get('screenwriter', {}).get('consensus', script)}")
        if "spatial" in dept_results:
            sp_label = "空间板块：" if is_zh else "Spatial:"
            parts.append(f"{sp_label}\n{dept_results['spatial']['consensus']}")
        if "storyboard" in dept_results:
            sb_label = "分镜部：" if is_zh else "Storyboard:"
            parts.append(f"{sb_label}\n{dept_results['storyboard']['consensus']}")
        return "\n\n---\n\n".join(parts)
    
    # P4: Sound + Editing
    elif dept_key == "sound":
        parts = []
        sw_label = "编剧部：" if is_zh else "Screenwriter:"
        parts.append(f"{sw_label}\n{dept_results.get('screenwriter', {}).get('consensus', script)}")
        if "storyboard" in dept_results:
            sb_label = "分镜部：" if is_zh else "Storyboard:"
            parts.append(f"{sb_label}\n{dept_results['storyboard']['consensus']}")
        return "\n\n---\n\n".join(parts)
    
    elif dept_key == "editing":
        parts = []
        for pk in ["screenwriter", "spatial", "storyboard", "dp", "lighting", "vfx", "sound"]:
            if pk in dept_results:
                p = DEPARTMENTS.get(pk, {})
                pn = (p.get("zh_name") if is_zh else p.get("en_name")) or p.get("zh_name") or p.get("en_name") or "?"
                consensus_label = "共识" if is_zh else "Consensus"
                parts.append(f"{pn}{consensus_label}：\n{dept_results[pk]['consensus']}" if is_zh else f"{pn} {consensus_label}:\n{dept_results[pk]['consensus']}")
        return "\n\n---\n\n".join(parts)
    
    else:
        # Academic/generic mode: always include research topic as context
        topic_label = "研究主题" if is_zh else "Research Topic"
        topic_line = f"{topic_label}：\n{script}" if is_zh else f"{topic_label}:\n{script}"
        # v0.11: Inject user-confirmed real papers for post-search academic departments
        _papers_summary = st.session_state.get("_papers_summary", "")
        if _papers_summary:
            _papers_label = ("已确认的真实文献列表（你的讨论必须基于这些真实文献展开，可引用其编号如[1][2]，不得虚构文献内容）" if is_zh
                             else "Confirmed real literature (your discussion MUST be grounded in these real papers; you may cite them as [1][2]; do NOT fabricate paper content)")
            topic_line += f"\n\n{_papers_label}:\n{_papers_summary}"
        prev = []
        try:
            _idx = DEPT_ORDER.index(dept_key)
        except ValueError:
            _idx = len(DEPT_ORDER)
        for pk in DEPT_ORDER[:_idx]:
            if pk in dept_results:
                p = DEPARTMENTS.get(pk, {})
                pn = p.get("zh_name" if is_zh else "en_name", pk)
                prev.append(f"{pn}：\n{dept_results[pk]['consensus']}" if is_zh else f"{pn}:\n{dept_results[pk]['consensus']}")
        if not prev:
            return topic_line
        return topic_line + "\n\n---\n\n" + "\n\n---\n\n".join(prev)


# ============ One-Click Full Debate ============

def run_all_debates(progress_callback=None):
    """Execute full debate flow (non-step mode).
    
    progress_callback: Optional callable(progress: float, text: str) for background thread mode.
    When provided, all progress updates go through this callback instead of st.progress().
    """
    is_zh = st.session_state.lang == "zh"
    api_url = st.session_state.api_url
    api_key = st.session_state.api_key
    model = st.session_state.model_name
    rounds = st.session_state.debate_rounds
    lang = st.session_state.lang
    extra = st.session_state.extra_instructions
    
    script = st.session_state.script
    positive = st.session_state.positive_prompt
    negative = st.session_state.negative_prompt
    chars = st.session_state.character_refs
    
    # v2.3: Token stats initialization
    stats = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "api_calls": 0}
    st.session_state.current_stats = stats
    st.session_state.current_start_time = time.time()
    
    arch_mode = st.session_state.architecture_mode
    
    # === Single Agent Baseline Mode ===
    if arch_mode == "single_agent":
        if progress_callback:
            progress_callback(0.0, t("debate_progress"))
        else:
            progress = st.progress(0, text=t("debate_progress"))
        result = run_single_agent(
            user_script=script,
            positive_prompt=positive,
            negative_prompt=negative,
            character_refs=chars,
            api_url=api_url, api_key=api_key, model=model,
            lang=lang, stats=stats,
        )
        st.session_state.final_output = result
        st.session_state.debate_completed = True
        st.session_state.current_end_time = time.time()
        autosave_result("normal_result", result)
        if progress_callback:
            progress_callback(1.0, "✅ " + ("单Agent基线生成完成" if is_zh else "Single agent baseline complete"))
        else:
            progress.progress(1.0, text="✅ " + ("单Agent基线生成完成" if is_zh else "Single agent baseline complete"))
        return
    
    # === Pipeline / Expert Pool Mode ===
    # v0.10: Resume support — config fingerprint for isolation
    import hashlib as _hl
    _fp = _hl.md5(f"{api_url}|{api_key}|{model}|{script}|{lang}|{rounds}".encode()).hexdigest()
    _saved_fp = st.session_state.get("_debate_fingerprint")
    if _saved_fp and _saved_fp != _fp:
        # Config changed (different user/API/topic) — discard stale partial results
        st.session_state.dept_results = {}
        st.session_state.cross_results = []
        st.session_state._debate_phase = "departments"
    elif not _saved_fp:
        st.session_state._debate_phase = "departments"
    st.session_state._debate_fingerprint = _fp

    # Read from session_state (survives Streamlit rerun)
    dept_results = st.session_state.get("dept_results", {})
    _resume_phase = st.session_state.get("_debate_phase", "departments")
    total_steps = len(DEPT_ORDER) * rounds * 3
    if progress_callback:
        progress_callback(0.0, t("debate_progress"))
    else:
        progress = st.progress(0, text=t("debate_progress"))
    # Calculate step offset from already-completed departments
    step = sum(
        len(DEPARTMENTS.get(d, {}).get("debaters", [])) * rounds
        for d in DEPT_ORDER if d in dept_results
    )
    
    # === v0.11: Academic flow — literature dept first, then real search, then pause for user review ===
    _cfg_flow = (st.session_state.get("workgroup_config") or get_current_config() or {})
    _LIT = "literature_search"
    _academic_flow = (
        arch_mode != "expert_pool"
        and _detect_pipeline_mode(_cfg_flow) == "academic"
        and _LIT in DEPT_ORDER
    )
    if _academic_flow and _resume_phase in ("departments", "search"):
        # Phase A: literature_search department formulates the search strategy
        if _LIT not in dept_results:
            dept = DEPARTMENTS.get(_LIT, {})
            dept_name = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or _LIT
            _lit_msg = f"📚 {dept_name} - " + ("正在制定检索策略..." if is_zh else "formulating search strategy...")
            if progress_callback:
                progress_callback(0.05, _lit_msg)
            else:
                progress.progress(0.05, text=_lit_msg)
            _lit_extra = (extra or "") + "\n" + (LIT_STRATEGY_INSTRUCTION_ZH if is_zh else LIT_STRATEGY_INSTRUCTION_EN)
            result = run_department_debate(
                department_key=_LIT,
                input_content=build_dept_input(_LIT),
                api_url=api_url, api_key=api_key, model=model,
                rounds=rounds, lang=lang, extra_instructions=_lit_extra,
                progress_callback=None,
                carry_forward=st.session_state.carry_forward,
                stats=stats,
            )
            dept_results[_LIT] = result
            st.session_state.dept_results = dept_results
            notify_stage_complete(dept_name)

        # 三重保证：提取JSON → 补考生成 → (搜索阶段 fallback generate_domain_config)
        _lit_consensus = dept_results.get(_LIT, {}).get("consensus", "")
        _strategy = _extract_search_strategy(_lit_consensus)
        if not _strategy:
            print("[v0.11] No valid JSON strategy in literature consensus, regenerating via dedicated LLM call...")
            _strategy = _regenerate_strategy_json(_lit_consensus, script, api_url, api_key, model)
        st.session_state._search_strategy = _strategy

        # Phase B: execute real literature search
        _search_msg = "🔍 " + ("正在检索真实文献（OpenAlex / Semantic Scholar / arXiv，约1-2分钟）..." if is_zh else "Searching real papers (OpenAlex/S2/arXiv, ~1-2 min)...")
        if progress_callback:
            progress_callback(0.45, _search_msg)
        else:
            progress.progress(0.45, text=_search_msg)
        _sr = run_academic_search(
            user_topic=script or st.session_state.get("config_user_input", ""),
            strategy=_strategy,
            api_url=api_url, api_key=api_key, model=model, lang=lang,
        )
        st.session_state._search_result = _sr
        st.session_state._debate_phase = "search_review"
        st.session_state._debate_running = False
        return  # Pause — user confirms in render_search_review_panel

    if arch_mode == "expert_pool":
        # Expert pool mode: scene analysis first, then streamlined debate
        expert_result = run_expert_pool_debate(
            user_script=script,
            positive_prompt=positive,
            character_refs=chars,
            api_url=api_url, api_key=api_key, model=model,
            rounds=rounds, lang=lang, extra_instructions=extra,
            carry_forward=st.session_state.carry_forward,
            progress_callback=None,  # 简化进度展示
            stats=stats,
        )
        dept_results = expert_result["dept_results"]
        st.session_state.expert_pool_result = expert_result
    else:
        # P1-P4: Department debates (Pipeline of Consensus)
        for dept_key in DEPT_ORDER:
            # v0.10: Skip already-completed departments (resume after rerun)
            if dept_key in dept_results:
                continue
            dept = DEPARTMENTS.get(dept_key, {})
            dept_name = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or "?"
            dept_input = build_dept_input(dept_key)
        
            def on_progress(dk, rn, total_r, debater):
                nonlocal step
                d = DEPARTMENTS.get(dk, {})
                dn = (d.get("zh_name") if is_zh else d.get("en_name")) or d.get("zh_name") or d.get("en_name") or "?"
                if debater == "consensus":
                    # Consensus generation phase
                    pct = min(step / total_steps, 0.99)
                    text = f"🔄 {dn} - 正在生成部门总结（可能需要2-5分钟）..."
                else:
                    step += 1
                    dn2 = d["debaters"][debater]["zh_name"] if is_zh else d["debaters"][debater]["en_name"]
                    pct = min(step / total_steps, 0.99)
                    text = f"{dn} 第{rn}/{total_r}轮 - {dn2}"
                if progress_callback:
                    progress_callback(pct, text)
                else:
                    progress.progress(pct, text)
        
            result = run_department_debate(
                department_key=dept_key,
                input_content=dept_input,
                api_url=api_url, api_key=api_key, model=model,
                rounds=rounds, lang=lang, extra_instructions=extra,
                progress_callback=on_progress,
                carry_forward=st.session_state.carry_forward,
                stats=stats,
            )
            dept_results[dept_key] = result
            st.session_state.dept_results = dept_results  # v0.10: Persist immediately for rerun resume
            notify_stage_complete(dept_name)
    
    # P2: Spatial planning cross-review
    spatial_consensus = dept_results.get("spatial", {}).get("consensus", "")
    if spatial_consensus:
        with st.spinner("🔍 " + ("空间板块交叉审核..." if is_zh else "Spatial review...")):
            spatial_review = run_spatial_review(
                spatial_consensus=spatial_consensus,
                reviewer_departments=["storyboard", "dp", "editing"],
                api_url=api_url, api_key=api_key, model=model, lang=lang,
                stats=stats,
            )
            # Replace original with revised spatial consensus
            dept_results["spatial"]["consensus"] = spatial_review["revised_consensus"]
            st.session_state.spatial_review_result = spatial_review
            notify_stage_complete("空间交叉审核" if is_zh else "Spatial Review")
    
    st.session_state.dept_results = dept_results
    st.session_state._debate_phase = "cross"
    
    # P5: Cross-debate (re-run if not yet reached summary phase; cross-debate is fast and safe to re-run)
    cross_results = []
    if _resume_phase not in ("summary", "done"):
        for cd in P5_CROSS_DEBATES:
            a_key, b_key = cd.get("side_a"), cd.get("side_b")
            if not a_key or not b_key:
                continue
            if a_key in dept_results and b_key in dept_results:
                cr = run_cross_debate(
                    cross_config=cd,
                    dept_a_consensus=dept_results[a_key]["consensus"],
                    dept_b_consensus=dept_results[b_key]["consensus"],
                    api_url=api_url, api_key=api_key, model=model, lang=lang,
                    stats=stats,
                )
                cross_results.append(cr)
        st.session_state.cross_results = cross_results
    
    st.session_state._debate_phase = "summary"
    
    # P6: Summary AI — branch on academic vs animation mode
    _cfg = (st.session_state.get("workgroup_config") or get_current_config() or {})
    _is_academic = _detect_pipeline_mode(_cfg) == "academic"

    # v0.10: Skip summary if already completed (resume from rerun)
    if st.session_state.get("final_output") and _resume_phase == "summary":
        final = st.session_state.final_output
    elif _is_academic:
        final = run_academic_summary(
            user_topic=script or st.session_state.get("config_user_input", ""),
            all_consensus={k: v["consensus"] for k, v in dept_results.items()},
            cross_results=cross_results,
            api_url=api_url, api_key=api_key, model=model, lang=lang,
            stats=stats,
            confirmed_papers=st.session_state.get("_confirmed_papers"),
            confirmed_preprints=st.session_state.get("_confirmed_preprints"),
        )
        notify_stage_complete("学术综述报告生成" if is_zh else "Academic Report")
    else:
        final = run_summary(
            all_consensus={k: v["consensus"] for k, v in dept_results.items()},
            cross_results=cross_results,
            user_script=script, positive_prompt=positive,
            negative_prompt=negative, character_refs=chars,
            api_url=api_url, api_key=api_key, model=model, lang=lang,
            stats=stats,
        )
        notify_stage_complete("分镜表+视频提示词生成" if is_zh else "Storyboard+Video Prompt")
    st.session_state.final_output = final
    st.session_state.debate_completed = True
    st.session_state.current_end_time = time.time()
    st.session_state._debate_phase = "done"
    autosave_result("normal_result", final)
    if progress_callback:
        progress_callback(1.0, "✅ " + ("全部辩论完成！" if is_zh else "All debates complete!"))
    else:
        progress.progress(1.0, text="✅ " + ("全部辩论完成！" if is_zh else "All debates complete!"))





# ============ Step-by-Step Mode ============

def step_run_round(correction: str = ""):
    """Step mode: execute current department current round debate"""
    # Bounds check: config may have changed, making step_dept_index stale
    if st.session_state.step_dept_index >= len(DEPT_ORDER):
        is_zh_local = st.session_state.get("lang", "zh") == "zh"
        st.session_state.step_dept_index = 0
        st.session_state.step_phase = "idle"
        st.warning("⚠️ " + ("部门配置已变更，步骤已重置" if is_zh_local else "Config changed, steps reset"))
        return
    dept_key = DEPT_ORDER[st.session_state.step_dept_index]
    dept_input = build_dept_input(dept_key)
    
    if correction:
        st.session_state.step_corrections.append(correction)
    
    extra = st.session_state.extra_instructions
    if st.session_state.step_corrections:
        all_corrections = "\n".join(f"导演纠正{i+1}：{c}" for i, c in enumerate(st.session_state.step_corrections))
        extra = f"{extra}\n{all_corrections}" if extra else all_corrections
    
    cf = st.session_state.carry_forward if st.session_state.step_round == 1 else ""
    
    # Ensure stats are initialized
    if st.session_state.current_stats is None:
        st.session_state.current_stats = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "api_calls": 0}
        st.session_state.current_start_time = time.time()
    
    round_log, all_arguments = run_department_round(
        department_key=dept_key,
        round_num=st.session_state.step_round,
        input_content=dept_input,
        previous_arguments=st.session_state.step_all_arguments,
        api_url=st.session_state.api_url,
        api_key=st.session_state.api_key,
        model=st.session_state.model_name,
        lang=st.session_state.lang,
        extra_instructions=extra,
        carry_forward=cf,
        stats=st.session_state.current_stats,
    )
    
    st.session_state.step_debate_log.extend(round_log)
    st.session_state.step_all_arguments = all_arguments
    st.session_state.step_phase = "round_done"


def step_generate_consensus():
    """Step mode: generate consensus for current department"""
    # Bounds check: config may have changed, making step_dept_index stale
    if st.session_state.step_dept_index >= len(DEPT_ORDER):
        is_zh_local = st.session_state.get("lang", "zh") == "zh"
        st.session_state.step_dept_index = 0
        st.session_state.step_phase = "idle"
        st.warning("⚠️ " + ("部门配置已变更，步骤已重置" if is_zh_local else "Config changed, steps reset"))
        return
    dept_key = DEPT_ORDER[st.session_state.step_dept_index]
    dept_input = build_dept_input(dept_key)
    
    # Ensure stats are initialized
    if st.session_state.current_stats is None:
        st.session_state.current_stats = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "api_calls": 0}
        st.session_state.current_start_time = time.time()
    
    consensus = run_department_consensus(
        department_key=dept_key,
        all_arguments=st.session_state.step_all_arguments,
        input_content=dept_input,
        api_url=st.session_state.api_url,
        api_key=st.session_state.api_key,
        model=st.session_state.model_name,
        lang=st.session_state.lang,
        extra_instructions=st.session_state.extra_instructions,
        rounds=st.session_state.step_round,
        stats=st.session_state.current_stats,
    )
    
    st.session_state.dept_results[dept_key] = {
        "department": dept_key,
        "debate_log": st.session_state.step_debate_log,
        "consensus": consensus,
    }
    st.session_state.step_phase = "dept_done"
    dept_name = DEPARTMENTS.get(dept_key, {}).get("zh_name" if st.session_state.lang == "zh" else "en_name", dept_key)
    notify_stage_complete(dept_name)


def step_next_dept():
    """Step mode: advance to next department"""
    st.session_state.step_dept_index += 1
    st.session_state.step_round = 1
    st.session_state.step_all_arguments = []
    st.session_state.step_debate_log = []
    st.session_state.step_corrections = []
    
    if st.session_state.step_dept_index >= len(DEPT_ORDER):
        st.session_state.step_phase = "all_dept_done"
    else:
        st.session_state.step_phase = "idle"


def step_run_spatial_review():
    """Step mode: execute P2 spatial cross-review"""
    dept_results = st.session_state.dept_results
    spatial_consensus = dept_results.get("spatial", {}).get("consensus", "")
    if spatial_consensus:
        spatial_review = run_spatial_review(
            spatial_consensus=spatial_consensus,
            reviewer_departments=["storyboard", "dp", "editing"],
            api_url=st.session_state.api_url,
            api_key=st.session_state.api_key,
            model=st.session_state.model_name,
            lang=st.session_state.lang,
            stats=st.session_state.get("current_stats"),
        )
        dept_results["spatial"]["consensus"] = spatial_review["revised_consensus"]
        st.session_state.spatial_review_result = spatial_review


def step_run_cross_debates():
    """Step mode: execute P5 cross-debate"""
    dept_results = st.session_state.dept_results
    cross_results = []
    
    # P2: Spatial review (if not done yet)
    if not st.session_state.get("spatial_review_result"):
        step_run_spatial_review()
    
    # P5: Cross-debate
    for cd in P5_CROSS_DEBATES:
        a_key, b_key = cd.get("side_a"), cd.get("side_b")
        if not a_key or not b_key:
            continue
        if a_key in dept_results and b_key in dept_results:
            cr = run_cross_debate(
                cross_config=cd,
                dept_a_consensus=dept_results[a_key]["consensus"],
                dept_b_consensus=dept_results[b_key]["consensus"],
                api_url=st.session_state.api_url,
                api_key=st.session_state.api_key,
                model=st.session_state.model_name,
                lang=st.session_state.lang,
                stats=st.session_state.get("current_stats"),
            )
            cross_results.append(cr)
    st.session_state.cross_results = cross_results
    st.session_state.step_phase = "cross_done"


def step_run_summary():
    """Step mode: generate final output"""
    # Ensure stats are initialized
    if st.session_state.current_stats is None:
        st.session_state.current_stats = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "api_calls": 0}
        st.session_state.current_start_time = time.time()

    dept_results = st.session_state.dept_results
    _is_zh = st.session_state.lang == "zh"
    _cfg = (st.session_state.get("workgroup_config") or get_current_config() or {})
    _is_academic = _detect_pipeline_mode(_cfg) == "academic"

    if _is_academic:
        final = run_academic_summary(
            user_topic=st.session_state.script or st.session_state.get("config_user_input", ""),
            all_consensus={k: v["consensus"] for k, v in dept_results.items()},
            cross_results=st.session_state.cross_results,
            api_url=st.session_state.api_url,
            api_key=st.session_state.api_key,
            model=st.session_state.model_name,
            lang=st.session_state.lang,
            stats=st.session_state.current_stats,
        )
    else:
        final = run_summary(
            all_consensus={k: v["consensus"] for k, v in dept_results.items()},
            cross_results=st.session_state.cross_results,
            user_script=st.session_state.script,
            positive_prompt=st.session_state.positive_prompt,
            negative_prompt=st.session_state.negative_prompt,
            character_refs=st.session_state.character_refs,
            api_url=st.session_state.api_url,
            api_key=st.session_state.api_key,
            model=st.session_state.model_name,
            lang=st.session_state.lang,
            stats=st.session_state.current_stats,
        )
    st.session_state.final_output = final
    st.session_state.debate_completed = True
    st.session_state.current_end_time = time.time()
    autosave_result("normal_result", final)
    st.session_state.step_phase = "completed"


# ============ Input Tab ============

def render_input_tab():
    is_zh = st.session_state.get("lang", "zh") == "zh"
    
    
    
    
    # Sync pending auto-fill content (from requirement research Phase 4)
    _pending = st.session_state.pop("_pending_script_fill", None)
    if _pending:
        st.session_state.script = _pending
        st.session_state.script_input = _pending
    
    current_config = get_current_config() or {}
    # Determine if current config is academic/programming (non-animation)
    dept_keys = list(current_config.get("departments", {}).keys())
    is_academic = _detect_pipeline_mode() == "academic"
    is_animation = any(k in dept_keys for k in ["screenwriter", "storyboard", "spatial", "dp", "lighting", "vfx", "sound", "editing"])
    
    if is_academic:
        st.subheader("📋 " + ("研究需求输入" if is_zh else "Research Requirements"))
        st.session_state.script = st.text_area(
            "📋 " + ("研究主题与需求" if is_zh else "Research Topic & Requirements"),
            value=st.session_state.script,
            key="script_input",
            height=200,
            placeholder="例如：\n研究主题：碳价预测\n方法论：机器学习\n时间范围：近5年\n质量标准：CSSCI及以上" if is_zh else "e.g.:\nTopic: Carbon price prediction\nMethodology: Machine Learning\nTime: Last 5 years\nQuality: CSSCI+",
            label_visibility="collapsed",
        )
        # Academic mode hides positive/negative prompts and character refs, shows research constraints instead
        with st.expander("🔬 " + ("研究约束与补充" if is_zh else "Research Constraints"), expanded=False):
            st.session_state.positive_prompt = st.text_area(
                "🔍 " + ("补充检索词（可选，用于辅助论文搜索）" if is_zh else "Extra Search Terms (optional)"),
                value=st.session_state.positive_prompt,
                height=80,
                placeholder="补充检索关键词，留空则自动生成..." if is_zh else "Extra search keywords, leave empty for auto-generation...",
                label_visibility="collapsed",
            )
            st.session_state.negative_prompt = st.text_area(
                "🚫 " + ("排除范围" if is_zh else "Exclusions"),
                value=st.session_state.negative_prompt,
                height=60,
                placeholder="排除的方向、不感兴趣的主题..." if is_zh else "Topics to exclude...",
                label_visibility="collapsed",
            )
            st.session_state.character_refs = ""  # 学术模式不需要角色引用
        st.info("💡 " + ("学术调研模式：输入研究主题和需求，各部门将围绕学术辩论展开" if is_zh else "Academic mode: Enter research topic, departments will debate academically"))
    else:
        # Generic input for non-academic modes
        _input_label = "📝 输入内容" if is_zh else "📝 Input Content"
        _input_hint = "描述你的需求、主题或场景..." if is_zh else "Describe your requirements, topic or scene..."
        if is_animation:
            _input_label = "🎬 基础剧本" if is_zh else "🎬 Script"
            _input_hint = "输入你的场景、角色、情节骨架" if is_zh else "Enter scene, characters, plot skeleton..."
        st.subheader(_input_label)
        st.session_state.script = st.text_area(
            _input_label,
            value=st.session_state.script,
            key="script_input",
            height=150,
            placeholder=_input_hint,
            label_visibility="collapsed",
        )
        
        # Animation-specific fields (only show for animation configs)
        if is_animation:
            col1, col2 = st.columns(2)
            with col1:
                st.subheader(t("positive_prompt_label"))
                st.session_state.positive_prompt = st.text_area(
                    t("positive_prompt_label"),
                    value=st.session_state.positive_prompt,
                    height=120,
                    placeholder=t("positive_prompt_hint"),
                    label_visibility="collapsed",
                )
            with col2:
                st.subheader(t("negative_prompt_label"))
                st.session_state.negative_prompt = st.text_area(
                    t("negative_prompt_label"),
                    value=st.session_state.negative_prompt,
                    height=120,
                    placeholder=t("negative_prompt_hint"),
                    label_visibility="collapsed",
                )
            st.subheader(t("character_refs_label"))
            st.session_state.character_refs = st.text_area(
                t("character_refs_label"),
                value=st.session_state.character_refs,
                height=100,
                placeholder=t("character_refs_hint"),
                label_visibility="collapsed",
            )
            st.info(t("segment_note"))
        else:
            # Generic mode: supplementary details in expander
            with st.expander("⚙️ " + ("补充设置" if is_zh else "Additional Settings"), expanded=False):
                st.session_state.positive_prompt = st.text_area(
                    "🔍 " + ("关键词/重点方向" if is_zh else "Keywords / Focus"),
                    value=st.session_state.positive_prompt,
                    height=80,
                    placeholder="补充关键词或重点..." if is_zh else "Additional keywords or focus...",
                )
                st.session_state.negative_prompt = st.text_area(
                    "🚫 " + ("排除项" if is_zh else "Exclusions"),
                    value=st.session_state.negative_prompt,
                    height=60,
                    placeholder="排除的内容..." if is_zh else "Items to exclude...",
                )
                st.session_state.character_refs = ""  # 非动画模式不需要角色参考
    
    # Carry-forward document input
    st.divider()
    st.subheader(t("carry_forward_label"))
    st.session_state.carry_forward = st.text_area(
        t("carry_forward_label"),
        value=st.session_state.carry_forward,
        height=120,
        placeholder=t("carry_forward_hint"),
        label_visibility="collapsed",
    )
    
    if not st.session_state.api_key:
        st.warning(t("need_api"))
        return
    
    if not st.session_state.script.strip():
        st.warning(t("need_script"))
        return
    
    # Two start buttons
    btn_col1, btn_col2 = st.columns(2)
    
    with btn_col1:
        debate_disabled = st.session_state.get("_debate_running", False)
        if st.button(t("start_debate"), type="secondary", use_container_width=True,
                     disabled=debate_disabled):
            st.session_state.step_phase = "idle"
            st.session_state.pipeline_mode = None  # will be auto-detected by _detect_pipeline_mode
            st.session_state.dept_results = {}
            st.session_state.cross_results = []
            st.session_state.final_output = {}
            st.session_state.debate_completed = False
            st.session_state.proofread_result = None
            st.session_state.spatial_review_result = None
            st.session_state._debate_running = True
            st.session_state._debate_fingerprint = None  # v0.10: clear for fresh start
            st.session_state._debate_phase = "departments"
            st.session_state._tab_index = 1  # switch to debate tab
            st.rerun()
    
    with btn_col2:
        if st.button(t("start_step"), type="primary", use_container_width=True):
            st.session_state.pipeline_mode = None  # will be auto-detected by _detect_pipeline_mode
            st.session_state.dept_results = {}
            st.session_state.cross_results = []
            st.session_state.final_output = {}
            st.session_state.debate_completed = False
            st.session_state.proofread_result = None
            st.session_state.spatial_review_result = None
            st.session_state.step_dept_index = 0
            st.session_state.step_round = 1
            st.session_state.step_all_arguments = []
            st.session_state.step_debate_log = []
            st.session_state.step_corrections = []
            st.session_state._debate_fingerprint = None  # v0.10
            st.session_state._debate_phase = "departments"  # v0.10
            step_run_round()
            st.rerun()


# ============ Search Review Panel (v0.11) ============

def render_search_review_panel():
    """学术模式 search_review 阶段：展示真实检索结果，用户确认后才进入部门辩论。"""
    is_zh = st.session_state.lang == "zh"
    sr = st.session_state.get("_search_result") or {}
    papers = sr.get("papers", []) or []
    preprints = sr.get("preprints", []) or []
    info = sr.get("search_info", {}) or {}

    st.markdown("### 🔍 " + ("文献检索结果确认" if is_zh else "Confirm Literature Search Results"))

    # 检索策略回显
    queries = info.get("queries", []) or []
    strategy_src = info.get("strategy_source", "none")
    _src_label = {
        "literature_dept": ("文献检索组制定" if is_zh else "by Literature Search dept"),
        "domain_config": ("自动生成（检索组策略不可用，已回退）" if is_zh else "auto-generated (fallback)"),
        "fallback": ("基础回退策略" if is_zh else "basic fallback"),
        "none": ("无" if is_zh else "none"),
    }.get(strategy_src, strategy_src)

    if queries:
        with st.expander(("🔎 检索策略（" + _src_label + "）") if is_zh else ("🔎 Search strategy (" + _src_label + ")"), expanded=False):
            for q in queries:
                st.markdown(f"- `{q}`")
            _excl = ((sr.get("domain_config") or {}).get("exclusion_signals") or [])
            if _excl:
                st.markdown(("**排除信号（含以下内容的文献会被剔除）：**" if is_zh else "**Exclusion signals (papers containing these are removed):**"))
                for _s in _excl:
                    st.markdown(f"- `{_s}`")

    # 统计
    level_counts = {}
    for p in papers:
        lv = (p.get("quality_level") or "C")
        level_counts[lv] = level_counts.get(lv, 0) + 1
    _lv_detail = " · ".join(f"{lv}:{c}" for lv, c in sorted(level_counts.items()))

    col_a, col_b, col_c = st.columns(3)
    col_a.metric(("期刊论文" if is_zh else "Journal papers"), len(papers))
    col_b.metric(("预印本" if is_zh else "Preprints"), len(preprints))
    col_c.metric(("总抓取" if is_zh else "Total fetched"), info.get("total_fetched", 0))
    if _lv_detail:
        st.caption(("质量分级 " if is_zh else "Quality tiers ") + _lv_detail)

    if not papers and not preprints:
        st.warning("⚠️ " + ("本次检索未找到符合条件的文献。建议调整关键词重新搜索；若直接确认，将按无论文流程生成报告（不含参考文献）。" if is_zh
                            else "No papers found. Consider adjusting keywords and re-searching; confirming directly will generate a report without references."))

    # 期刊论文 checkbox 列表（默认全选）
    if papers:
        st.markdown("**" + ("勾选纳入后续分析的文献（默认全选）：" if is_zh else "Select papers for subsequent analysis (all selected by default):") + "**")
        for i, p in enumerate(papers):
            _lv = p.get("quality_level") or "?"
            _title = p.get("title") or "?"
            _journal = p.get("journal") or "?"
            _year = p.get("year") or "?"
            _cited = p.get("citation_count") or 0
            label = f"[{_lv}] {_title} — {_journal} ({_year}), " + (f"被引 {_cited}" if is_zh else f"cited {_cited}")
            st.checkbox(label, value=True, key=f"_sr_paper_{i}")

    # 预印本默认不勾选
    if preprints:
        _pre_title = (f"📄 预印本（未经同行评审，默认不纳入，共 {len(preprints)} 篇）" if is_zh
                      else f"📄 Preprints (not peer-reviewed, excluded by default, {len(preprints)})")
        with st.expander(_pre_title, expanded=False):
            for j, p in enumerate(preprints[:15]):
                _title = p.get("title") or "?"
                _year = p.get("year") or "?"
                st.checkbox(f"{_title} ({_year})", value=False, key=f"_sr_pre_{j}")

    st.divider()

    # 操作区
    btn_confirm, btn_research = st.columns(2)
    with btn_confirm:
        if st.button("✅ " + ("确认，开始部门辩论" if is_zh else "Confirm & start department debates"),
                     type="primary", use_container_width=True):
            selected = [p for i, p in enumerate(papers) if st.session_state.get(f"_sr_paper_{i}", True)]
            selected_pre = [p for j, p in enumerate(preprints[:15]) if st.session_state.get(f"_sr_pre_{j}", False)]
            st.session_state._confirmed_papers = selected
            st.session_state._confirmed_preprints = selected_pre
            st.session_state._papers_summary = build_papers_summary(selected + selected_pre)
            st.session_state._debate_phase = "dept_rest"
            st.session_state._debate_running = True
            st.rerun()
    with btn_research:
        with st.expander("🔄 " + ("调整关键词重新搜索" if is_zh else "Adjust keywords & re-search")):
            st.caption(("每行一个检索式（仅英文；中文检索式会被自动过滤）" if is_zh else "One query per line (English only; Chinese queries are auto-filtered)"))
            edited = st.text_area(
                "queries",
                value="\n".join(queries),
                height=150,
                label_visibility="collapsed",
                key="_sr_edit_queries",
            )
            if st.button(("执行重新搜索" if is_zh else "Run search again"), type="secondary", use_container_width=True):
                _is_zh_char = lambda s: any('\u4e00' <= c <= '\u9fff' for c in s)
                new_queries = [l.strip() for l in edited.split("\n") if l.strip() and not _is_zh_char(l.strip())]
                if new_queries:
                    # 清掉旧的勾选状态，避免新旧列表错位
                    for k in list(st.session_state.keys()):
                        if k.startswith("_sr_paper_") or k.startswith("_sr_pre_"):
                            del st.session_state[k]
                    _old_excl = (st.session_state.get("_search_strategy") or {}).get("exclusion_signals", [])
                    st.session_state._search_strategy = {"search_queries": new_queries, "exclusion_signals": _old_excl}
                    st.session_state._search_result = None
                    st.session_state._debate_phase = "search"
                    st.session_state._debate_running = True
                    st.rerun()
                else:
                    st.error("❌ " + ("没有可用的英文检索式" if is_zh else "No valid English queries"))


# ============ Debate Tab ============

def render_debate_tab():
    is_zh = st.session_state.lang == "zh"
    
    # Show persisted error from previous run (e.g. 401 Unauthorized) before it gets wiped by rerun
    if st.session_state.get("_debate_error"):
        st.error(f"❌ {st.session_state._debate_error}")
        del st.session_state._debate_error
    
    # v0.10: Resume support — _debate_running stays True during execution.
    # If rerun happens mid-exec, guard re-enters run_all_debates() which
    # reads dept_results from session_state and skips completed departments.
    # Config fingerprint ensures isolation (different user/API/topic = fresh start).
    if st.session_state.get("_debate_running"):
        try:
            run_all_debates()
        except Exception as e:
            st.session_state._debate_error = f"{'辩论执行出错' if is_zh else 'Debate execution error'}: {e}"
        finally:
            st.session_state._debate_running = False
            st.session_state._tab_index = 1  # ensure results tab after completion
        st.rerun()

    # v0.11: search_review phase — show paper confirmation panel, wait for user
    if (st.session_state.get("_debate_phase") == "search_review"
            and not st.session_state.get("debate_completed")):
        render_search_review_panel()
        return

    if st.session_state.step_mode and st.session_state.step_phase != "idle":
        render_step_mode()
        return
    
    dept_results = st.session_state.get("dept_results", {})
    
    if not dept_results:
        st.info(t("debate_not_started"))
        return
    
    for dept_key in DEPT_ORDER:
        if dept_key not in dept_results:
            continue
        
        dept = DEPARTMENTS.get(dept_key, {})
        result = dept_results[dept_key]
        dept_name = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or "?"
        
        with st.expander(f"🎬 **{dept_name}**", expanded=False):
            for log_entry in result.get("debate_log", []):
                st.markdown(f"**{log_entry['debater_name']}** (Round {log_entry['round']})")
                st.markdown(log_entry["content"])
                st.divider()
            
            st.subheader(t("consensus"))
            st.markdown(result.get("consensus", ""))
            
            st.markdown("---")
            st.markdown("**" + t("director_review") + "**")
            rejection = st.text_input(
                t("rejection_reason"),
                key=f"reject_{dept_key}",
                placeholder=t("reject_hint"),
                label_visibility="collapsed",
            )
            col_a, col_b = st.columns(2)
            with col_a:
                st.button(t("approve"), key=f"approve_{dept_key}", disabled=True, use_container_width=True)
            with col_b:
                if st.button(t("reject_redebate"), key=f"reject_btn_{dept_key}", use_container_width=True):
                    if rejection.strip():
                        rerun_single_dept(dept_key, rejection)
                    else:
                        st.warning(t("reject_hint"))


def render_step_mode():
    """Render step-by-step mode UI"""
    is_zh = st.session_state.lang == "zh"
    phase = st.session_state.step_phase
    dept_index = st.session_state.step_dept_index
    
    if dept_index < len(DEPT_ORDER):
        current_dept = DEPARTMENTS.get(DEPT_ORDER[dept_index] if dept_index < len(DEPT_ORDER) else "", {})
        dept_name = (current_dept.get("zh_name") if is_zh else current_dept.get("en_name")) or current_dept.get("zh_name") or current_dept.get("en_name") or "?"
    else:
        dept_name = "---"
    
    st.markdown(f"**{t('step_current_status')}：** {dept_name} · {t('step_round_title', round=st.session_state.step_round)} · `{phase}`")
    
    if phase == "round_done":
        dept_key = DEPT_ORDER[dept_index]
        dept = DEPARTMENTS.get(dept_key, {})
        dept_name = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or "?"
        
        st.subheader(t("step_round_title", round=st.session_state.step_round) + f" — {dept_name}")
        
        round_num = st.session_state.step_round
        for log_entry in st.session_state.step_debate_log:
            if log_entry["round"] == round_num:
                with st.chat_message("assistant"):
                    st.markdown(f"**{log_entry['debater_name']}**")
                    st.markdown(log_entry["content"])
        
        st.divider()
        
        correction = st.text_input(
            t("step_correction_label"),
            key=f"step_correction_{dept_index}_{round_num}",
            placeholder=t("step_correction_hint"),
        )
        
        col1, col2 = st.columns(2)
        max_rounds = st.session_state.debate_rounds
        is_last_round = st.session_state.step_round >= max_rounds
        
        with col1:
            if is_last_round:
                if st.button(t("step_finish_dept"), type="primary", use_container_width=True, key=f"step_finish_{dept_index}"):
                    with st.spinner("🔄 " + ("正在生成部门总结（可能需要2-5分钟）..." if is_zh else "Generating consensus (may take 2-5 min)...")):
                        step_generate_consensus()
                    st.rerun()
            else:
                if st.button(t("step_continue"), type="primary", use_container_width=True, key=f"step_continue_{dept_index}_{round_num}"):
                    st.session_state.step_round += 1
                    step_run_round(correction=correction)
                    st.rerun()
        
        with col2:
            if not is_last_round:
                if st.button(t("step_finish_dept"), use_container_width=True, key=f"step_early_finish_{dept_index}_{round_num}"):
                    with st.spinner("🔄 " + ("正在生成部门总结（可能需要2-5分钟）..." if is_zh else "Generating consensus (may take 2-5 min)...")):
                        step_generate_consensus()
                    st.rerun()
    
    elif phase == "dept_done":
        dept_key = DEPT_ORDER[dept_index]
        dept = DEPARTMENTS.get(dept_key, {})
        dept_name = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or "?"
        result = st.session_state.dept_results.get(dept_key, {})
        
        st.subheader(t("step_dept_done", dept=dept_name))
        
        with st.expander("📝 " + ("辩论过程" if is_zh else "Debate Log"), expanded=False):
            for log_entry in result.get("debate_log", []):
                st.markdown(f"**{log_entry['debater_name']}** (Round {log_entry['round']})")
                st.markdown(log_entry["content"])
                st.divider()
        
        st.subheader(t("consensus"))
        st.markdown(result.get("consensus", ""))
        
        st.markdown("---")
        st.markdown("**" + t("director_review") + "**")
        rejection = st.text_input(
            t("rejection_reason"),
            key=f"step_reject_{dept_key}",
            placeholder=t("reject_hint"),
            label_visibility="collapsed",
        )
        col_a, col_b = st.columns(2)
        with col_a:
            if st.button(t("approve"), type="primary", use_container_width=True, key=f"step_approve_{dept_key}"):
                step_next_dept()
                if st.session_state.step_phase == "all_dept_done":
                    st.rerun()
                else:
                    step_run_round()
                    st.rerun()
        with col_b:
            if st.button(t("reject_redebate"), use_container_width=True, key=f"step_reject_btn_{dept_key}"):
                if rejection.strip():
                    st.session_state.step_round = 1
                    st.session_state.step_all_arguments = []
                    st.session_state.step_debate_log = []
                    st.session_state.step_corrections = [rejection]
                    step_run_round()
                    st.rerun()
                else:
                    st.warning(t("reject_hint"))
    
    elif phase == "all_dept_done":
        st.success(t("step_all_done"))
        
        for dept_key in DEPT_ORDER:
            if dept_key in st.session_state.dept_results:
                dept = DEPARTMENTS.get(dept_key, {})
                dept_name = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or "?"
                result = st.session_state.dept_results[dept_key]
                with st.expander(f"📋 **{dept_name}**"):
                    st.markdown(result.get("consensus", ""))
        
        if st.button(t("step_run_cross"), type="primary", use_container_width=True, key="step_run_cross"):
            with st.spinner("⚔️ " + ("交叉辩论中..." if is_zh else "Cross-debating...")):
                step_run_cross_debates()
            st.rerun()
    
    elif phase == "cross_done":
        st.success("✅ " + ("交叉辩论完成" if is_zh else "Cross-debates complete"))
        
        # P2 spatial review result
        if st.session_state.get("spatial_review_result"):
            with st.expander(f"🔍 **{t('p2_cross_title')}**"):
                sr = st.session_state.spatial_review_result
                for dk, rv in sr.get("reviews", {}).items():
                    dept = DEPARTMENTS.get(dk, {})
                    dept_name = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or "?"
                    st.markdown(f"**{dept_name}** Review feedback:")
                    st.markdown(rv)
                    st.divider()
        
        # P5 cross-debate result
        for cr in st.session_state.cross_results:
            a_dept = DEPARTMENTS.get(cr.get("side_a", ""), {})
            b_dept = DEPARTMENTS.get(cr.get("side_b", ""), {})
            a_name = a_dept.get("zh_name" if is_zh else "en_name", cr.get("side_a", "?"))
            b_name = b_dept.get("zh_name" if is_zh else "en_name", cr.get("side_b", "?"))
            with st.expander(f"⚔️ **{a_name} vs {b_name}** — {cr['topic']}"):
                st.markdown(cr.get("debate_result", ""))
        
        if st.button(t("step_run_summary"), type="primary", use_container_width=True, key="step_run_summary"):
            with st.spinner("🎬 " + ("生成最终产出中..." if is_zh else "Generating final output...")):
                step_run_summary()
                _mode_label = "学术综述" if st.session_state.get("pipeline_mode") == "academic" else "动画分镜"
                st.toast(f"📋 {'产出模式' if is_zh else 'Output mode'}: {_mode_label}", icon="📋")
            st.rerun()
    
    elif phase == "completed":
        st.success("🎬 " + ("全部流程完成！请到「最终产出」Tab查看结果，到「校对」Tab执行校对" if is_zh else "All complete! Check Output tab for results, Proofread tab for QA"))
    
    # Existing department results
    if st.session_state.dept_results and phase not in ("all_dept_done", "cross_done", "completed"):
        st.divider()
        for dk in DEPT_ORDER:
            if dk in st.session_state.dept_results and DEPT_ORDER.index(dk) < dept_index:
                dept = DEPARTMENTS.get(dk, {})
                dept_name = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or "?"
                result = st.session_state.dept_results[dk]
                with st.expander(f"✅ **{dept_name}**"):
                    st.markdown(result.get("consensus", ""))


def rerun_single_dept(dept_key: str, revision_note: str):
    """Reject and re-debate a single department"""
    dept_input = build_dept_input(dept_key)
    extra = st.session_state.extra_instructions
    if revision_note:
        extra = f"{extra}\n导演修改意见：{revision_note}" if extra else f"导演修改意见：{revision_note}"
    
    result = run_department_debate(
        department_key=dept_key,
        input_content=dept_input,
        api_url=st.session_state.api_url,
        api_key=st.session_state.api_key,
        model=st.session_state.model_name,
        rounds=st.session_state.debate_rounds,
        lang=st.session_state.lang,
        extra_instructions=extra,
        carry_forward=st.session_state.carry_forward,
        stats=st.session_state.get("current_stats"),
    )
    st.session_state.dept_results[dept_key] = result
    st.rerun()


# ============ Cross-Debate Tab ============

def render_cross_tab():
    
    
    
    cross_results = st.session_state.get("cross_results", [])
    is_zh = st.session_state.lang == "zh"
    
    # P2 spatial review
    if st.session_state.get("spatial_review_result"):
        sr = st.session_state.spatial_review_result
        st.subheader(t("p2_cross_title"))
        for dk, rv in sr.get("reviews", {}).items():
            dept = DEPARTMENTS.get(dk, {})
            dept_name = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or "?"
            with st.expander(f"🔍 **{dept_name}** " + ("审核反馈" if is_zh else "Review")):
                st.markdown(rv)
        with st.expander("📝 " + ("空间板块修订方案" if is_zh else "Revised Spatial Plan")):
            st.markdown(sr.get("revised_consensus", ""))
    
    st.divider()
    
    # P5 cross-debate
    st.subheader(t("p5_cross_title"))
    if not cross_results:
        st.info(t("cross_not_started"))
        return
    
    for cr in cross_results:
        a_dept = DEPARTMENTS.get(cr.get("side_a", ""), {})
        b_dept = DEPARTMENTS.get(cr.get("side_b", ""), {})
        a_name = (a_dept.get("zh_name") if is_zh else a_dept.get("en_name")) or a_dept.get("zh_name") or a_dept.get("en_name") or cr.get("side_a", "?")
        b_name = (b_dept.get("zh_name") if is_zh else b_dept.get("en_name")) or b_dept.get("zh_name") or b_dept.get("en_name") or cr.get("side_b", "?")
        
        with st.expander(f"⚔️ **{a_name} vs {b_name}** — {cr['topic']}", expanded=False):
            st.markdown(cr.get("debate_result", ""))


# ============ Output Tab ============

def render_output_tab():
    is_zh = st.session_state.lang == "zh"
    
    
    
    
    final = st.session_state.get("final_output", {})
    
    # Detect mode
    _cfg = (st.session_state.get("workgroup_config") or get_current_config() or {})
    is_academic = _detect_pipeline_mode(_cfg) == "academic"
    is_animation = not is_academic
    
    # 🔧 Auto-restore normal mode results
    if not final:
        saved = autosave_load("normal_result")
        if saved:
            st.info("💾 " + ("检测到上次辩论结果，已自动恢复！" if is_zh else "Previous debate result recovered from disk!"))
            st.session_state.final_output = saved
            st.session_state.debate_completed = True
            final = saved
    
    if not final:
        st.info(t("output_not_ready"))
        return
    
    if is_academic:
        _render_academic_output(final, is_zh)
    else:
        _render_animation_output(final, is_zh)


def _render_academic_output(final, is_zh):
    """Academic mode output: research report + department consensus + cross-debate + Word export"""
    dept_results = st.session_state.get("dept_results", {})
    cross_results = st.session_state.get("cross_results", [])
    
    # ---- Search Info (if available) ----
    search_info = final.get("search_info") if isinstance(final, dict) else None
    if search_info:
        with st.expander("🔍 " + ("文献检索信息" if is_zh else "Literature Search Info"), expanded=False):
            if is_zh:
                st.write(f"**检索词：** {search_info.get('query', 'N/A')}")
                st.write(f"**总抓取：** {search_info.get('total_fetched', 0)}篇 | "
                         f"**过滤后：** {search_info.get('after_filter', 0)}篇期刊论文 | "
                         f"**预印本：** {search_info.get('preprint_count', 0)}篇")
                st.write(f"**参考文献数：** {search_info.get('papers_in_refs', 0)}篇 | "
                         f"**含真实引用：** {'是' if search_info.get('has_real_references') else '否'}")
                _dc_active = search_info.get('domain_config_active', False)
                _q_count = search_info.get('queries_used', 0)
                st.write(f"**搜索轮次：** {_q_count}轮 | "
                         f"**智能领域配置：** {'已启用' if _dc_active else '未启用(回退模式)'}")
                if search_info.get('after_filter', 0) == 0 and search_info.get('preprint_count', 0) > 0:
                    st.warning("⚠️ 期刊论文搜索结果为0，仅有arXiv预印本。可能原因：Semantic Scholar/OpenAlex API网络不通，或检索词匹配度低。")
            else:
                st.write(f"**Query:** {search_info.get('query', 'N/A')}")
                st.write(f"**Total fetched:** {search_info.get('total_fetched', 0)} | "
                         f"**After filter:** {search_info.get('after_filter', 0)} journal | "
                         f"**Preprints:** {search_info.get('preprint_count', 0)}")
                _dc_active = search_info.get('domain_config_active', False)
                _q_count = search_info.get('queries_used', 0)
                st.write(f"**Queries:** {_q_count} rounds | "
                         f"**Domain config:** {'Active' if _dc_active else 'Fallback mode'}")
                if search_info.get('after_filter', 0) == 0 and search_info.get('preprint_count', 0) > 0:
                    st.warning("⚠️ No journal papers found, only arXiv preprints. Possible: Semantic Scholar/OpenAlex API unreachable.")
    
    # ---- Final Research Report ----
    st.subheader("📋 " + ("最终研究报告" if is_zh else "Final Research Report"))
    
    # Get the final report text — prioritize the academic summary output
    report_text = final.get("final_report", "") or final.get("consensus_report", "") or ""
    if not report_text and dept_results:
        # Fallback: assemble from department consensus
        parts = []
        for dept_key, dept_data in dept_results.items():
            name = dept_data.get("zh_name", dept_key) if is_zh else dept_data.get("en_name", dept_key)
            consensus = dept_data.get("consensus", "")
            if consensus:
                parts.append(f"## {name}\n\n{consensus}")
        if parts:
            report_text = "\n\n---\n\n".join(parts)
    
    if report_text:
        st.markdown(report_text)
        
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            st.download_button(
                "📥 " + ("下载研究报告 (MD)" if is_zh else "Download Report (MD)"),
                data=report_text,
                file_name="research_report.md",
                mime="text/markdown",
                use_container_width=True,
                key="dl_academic_report_md",
            )
        with col_dl2:
            # Word export using python-docx
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from io import BytesIO
                
                doc = Document()
                # Parse markdown into docx paragraphs (basic conversion)
                for line in report_text.split("\n"):
                    line_stripped = line.strip()
                    if not line_stripped:
                        continue
                    if line_stripped.startswith("# "):
                        doc.add_heading(line_stripped[2:], level=1)
                    elif line_stripped.startswith("## "):
                        doc.add_heading(line_stripped[3:], level=2)
                    elif line_stripped.startswith("### "):
                        doc.add_heading(line_stripped[4:], level=3)
                    elif line_stripped.startswith("#### "):
                        doc.add_heading(line_stripped[5:], level=4)
                    elif line_stripped.startswith("- ") or line_stripped.startswith("* "):
                        doc.add_paragraph(line_stripped[2:], style="List Bullet")
                    elif line_stripped.startswith("---"):
                        doc.add_paragraph("─" * 40)
                    else:
                        # Handle bold markers with state toggle
                        p = doc.add_paragraph()
                        parts = line_stripped.split("**")
                        _bold = False
                        for part in parts:
                            if part:
                                run = p.add_run(part)
                                if _bold:
                                    run.bold = True
                            _bold = not _bold
                
                buf = BytesIO()
                doc.save(buf)
                docx_bytes = buf.getvalue()
                
                st.download_button(
                    "📝 " + ("下载研究报告 (Word)" if is_zh else "Download Report (Word)"),
                    data=docx_bytes,
                    file_name="research_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="dl_academic_report_docx",
                )
            except ImportError:
                st.caption("💡 " + ("安装 python-docx 可导出 Word: pip install python-docx" if is_zh else "Install python-docx for Word export: pip install python-docx"))
            except Exception as _e:
                st.warning(f"Word export failed: {_e}")
    else:
        st.info("📝 " + ("暂无最终报告，请先完成辩论流程" if is_zh else "No final report yet, complete debate first"))
    
    st.divider()
    
    # ---- Programming & Tutorial Outputs (Priority Panel) ----
    _prog = dept_results.get("programming", {})
    _tut = dept_results.get("tutorial", {})
    _prog_consensus = _prog.get("consensus", "") if isinstance(_prog, dict) else ""
    _tut_consensus = _tut.get("consensus", "") if isinstance(_tut, dict) else ""
    
    if _prog_consensus or _tut_consensus:
        st.subheader("💻 " + ("代码实现与教程" if is_zh else "Code & Tutorial"))
        col_code, col_tut = st.columns(2)
        with col_code:
            if _prog_consensus:
                with st.container(border=True):
                    st.markdown("### 🐍 " + ("代码实现" if is_zh else "Code Implementation"))
                    st.markdown(_prog_consensus)
                    st.download_button(
                        "📥 " + ("下载代码产出 (MD)" if is_zh else "Download Code (MD)"),
                        data=_prog_consensus,
                        file_name="programming_output.md",
                        mime="text/markdown",
                        use_container_width=True,
                        key="dl_programming_output",
                    )
        with col_tut:
            if _tut_consensus:
                with st.container(border=True):
                    st.markdown("### 📖 " + ("教程" if is_zh else "Tutorial"))
                    st.markdown(_tut_consensus)
                    st.download_button(
                        "📥 " + ("下载教程产出 (MD)" if is_zh else "Download Tutorial (MD)"),
                        data=_tut_consensus,
                        file_name="tutorial_output.md",
                        mime="text/markdown",
                        use_container_width=True,
                        key="dl_tutorial_output",
                    )
    
    # ---- Department Consensus Details ----
    _academic_label_depts = {"programming", "tutorial"}
    if dept_results:
        st.subheader("🏢 " + ("各部门共识详情" if is_zh else "Department Consensus Details"))
        _non_empty = False
        for dept_key, dept_data in dept_results.items():
            name = dept_data.get("zh_name", dept_key) if is_zh else dept_data.get("en_name", dept_key)
            consensus = dept_data.get("consensus", "")
            if consensus:
                _non_empty = True
                # Add type label for programming/tutorial departments
                _label = ""
                if dept_key == "programming":
                    _label = " [💻 代码产出]" if is_zh else " [💻 Code Output]"
                elif dept_key == "tutorial":
                    _label = " [📖 教程产出]" if is_zh else " [📖 Tutorial Output]"
                with st.expander(f"🏢 {name}{_label}", expanded=False):
                    st.markdown(consensus)
        if not _non_empty:
            st.info("📝 " + ("暂无部门共识，请先完成辩论流程" if is_zh else "No department consensus yet"))
    
    st.divider()
    
    # ---- Paper Metadata CSV Export ----
    _papers_csv = st.session_state.get("papers_metadata_csv")
    if not _papers_csv:
        # Try to find papers_metadata.csv from disk
        import os as _os
        _csv_paths = [
            _os.path.join(_os.getcwd(), "papers_metadata.csv"),
            _os.path.join(_os.getcwd(), "output", "papers_metadata.csv"),
        ]
        for _cp in _csv_paths:
            if _os.path.isfile(_cp):
                try:
                    with open(_cp, "r", encoding="utf-8") as _f:
                        _papers_csv = _f.read()
                    st.session_state.papers_metadata_csv = _papers_csv
                    break
                except Exception:
                    pass
    
    if _papers_csv:
        st.subheader("📊 " + ("论文元数据导出" if is_zh else "Paper Metadata Export"))
        col_csv1, col_csv2 = st.columns([3, 1])
        with col_csv1:
            st.caption(f"📄 papers_metadata.csv ({len(_papers_csv.split(chr(10)))-1} 行)" if is_zh else f"📄 papers_metadata.csv ({len(_papers_csv.split(chr(10)))-1} rows)")
        with col_csv2:
            st.download_button(
                "📥 CSV",
                data=_papers_csv,
                file_name="papers_metadata.csv",
                mime="text/csv",
                use_container_width=True,
                key="dl_papers_metadata",
            )
    
    # ---- Chart Export ----
    import os as _os2
    _chart_dir = _os2.path.join(_os2.getcwd(), "charts")
    if _os2.path.isdir(_chart_dir):
        _chart_files = sorted([f for f in _os2.listdir(_chart_dir) if f.endswith((".png", ".jpg", ".svg", ".html"))])
        if _chart_files:
            st.subheader("📈 " + ("可视化图表" if is_zh else "Charts"))
            _chart_cols = st.columns(min(len(_chart_files), 3))
            for _i, _cf in enumerate(_chart_files):
                _fp = _os2.path.join(_chart_dir, _cf)
                with _chart_cols[_i % 3]:
                    if _cf.endswith(".html"):
                        with open(_fp, "r", encoding="utf-8") as _hf:
                            st.components.v1.html(_hf.read(), height=300, scrolling=True)
                    else:
                        st.image(_fp, caption=_cf, use_container_width=True)
                    with open(_fp, "rb") as _bf:
                        st.download_button(
                            f"📥 {_cf}",
                            data=_bf.read(),
                            file_name=_safe_filename(_cf),
                            mime="image/png" if _cf.endswith(".png") else "image/svg+xml",
                            use_container_width=True,
                            key=f"dl_chart_{_i}",
                        )
    
    # ---- Cross-Debate Results ----
    if cross_results:
        st.subheader("🔀 " + ("交叉辩论结果" if is_zh else "Cross-Debate Results"))
        for i, cr in enumerate(cross_results):
            cr_title = cr.get("title", f"Cross-debate {i+1}") if isinstance(cr, dict) else f"Cross-debate {i+1}"
            cr_content = cr.get("result", "") or cr.get("consensus", "") if isinstance(cr, dict) else str(cr)
            if cr_content:
                with st.expander(f"🔀 {cr_title}", expanded=False):
                    st.markdown(cr_content)

def _render_animation_output(final, is_zh):
    """Animation mode output: storyboard + video prompt + spatial diagram + asset checklist + carry-forward + re-roll"""
    # Storyboard
    st.subheader(t("output_storyboard"))
    storyboard_text = final.get("storyboard_prompt", "")
    st.code(storyboard_text, language=None)
    
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            t("export_storyboard"),
            data=storyboard_text,
            file_name="storyboard_prompt.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with col2:
        st.download_button(
            t("copy_prompt"),
            data=storyboard_text,
            file_name="storyboard_prompt.txt",
            mime="text/plain",
            use_container_width=True,
        )
    
    st.divider()
    
    # Video prompts
    st.subheader(t("output_video"))
    video_text = final.get("video_prompt", "")
    st.code(video_text, language=None)
    
    col3, col4 = st.columns(2)
    with col3:
        st.download_button(
            t("export_video"),
            data=video_text,
            file_name="video_prompt.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with col4:
        st.download_button(
            t("copy_prompt"),
            data=video_text,
            file_name="video_prompt.txt",
            mime="text/plain",
            use_container_width=True,
        )
    
    st.divider()
    
    # Spatial diagram
    st.subheader(t("spatial_diagram"))
    st.caption(t("spatial_diagram_hint"))
    
    dept_results = st.session_state.get("dept_results", {})
    spatial_consensus = dept_results.get("spatial", {}).get("consensus", "") if dept_results else ""
    
    if not spatial_consensus:
        st.warning(t("spatial_diagram_need_spatial"))
    else:
        if st.button(t("spatial_diagram_generate"), type="primary", use_container_width=True, key="gen_spatial_diagram"):
            with st.spinner(t("spatial_diagram_generating")):
                diagram_result = run_spatial_diagram(
                    spatial_consensus=spatial_consensus,
                    user_script=st.session_state.script,
                    positive_prompt=st.session_state.positive_prompt,
                    api_url=st.session_state.api_url,
                    api_key=st.session_state.api_key,
                    model=st.session_state.model_name,
                    lang=st.session_state.lang,
                    stats=st.session_state.get("current_stats"),
                )
                st.session_state.spatial_diagram_result = diagram_result
        
        diagram = st.session_state.get("spatial_diagram_result")
        if diagram:
            scene_count = diagram.get("scene_count", 1)
            scene_diagrams = diagram.get("scene_diagrams", [])
            
            if scene_count > 1 and scene_diagrams:
                for sd in scene_diagrams:
                    st.markdown(f"**🗺️ {sd['scene_name']}**")
                    st.code(sd["diagram_prompt"], language=None)
                    st.download_button(
                        f"📥 {t('spatial_diagram_export')} - {sd['scene_name']}",
                        data=sd["diagram_prompt"],
                        file_name=_safe_filename(f"spatial_diagram_{sd['scene_name']}.txt"),
                        mime="text/plain",
                        use_container_width=True,
                        key=f"dl_spatial_{sd['scene_name']}",
                    )
                    st.divider()
            else:
                diagram_text = diagram.get("spatial_diagram_prompt", "")
                st.code(diagram_text, language=None)
                
                sd_col1, sd_col2 = st.columns(2)
                with sd_col1:
                    st.download_button(
                        t("spatial_diagram_export"),
                        data=diagram_text,
                        file_name="spatial_diagram_prompt.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )
                with sd_col2:
                    st.download_button(
                        t("copy_prompt"),
                        data=diagram_text,
                        file_name="spatial_diagram_prompt.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )
    
    st.divider()
    
    # Asset reference table
    st.subheader(t("asset_checklist"))
    st.caption(t("asset_checklist_hint"))
    
    dept_results = st.session_state.get("dept_results", {})
    if dept_results:
        asset_checklist_text = generate_asset_checklist(
            all_consensus={k: v["consensus"] for k, v in dept_results.items()},
            user_script=st.session_state.script,
            character_refs=st.session_state.character_refs,
            lang=st.session_state.lang,
        )
        st.markdown(asset_checklist_text)
        st.download_button(
            t("asset_checklist_download"),
            data=asset_checklist_text,
            file_name="asset_checklist.md",
            mime="text/markdown",
            use_container_width=True,
        )
    else:
        st.info(t("output_not_ready"))
    
    st.divider()
    
    
    
    # Carry-forward document
    st.subheader(t("carry_forward_label"))
    st.caption(t("carry_forward_note"))
    
    if st.button(t("carry_forward_generate"), type="primary", use_container_width=True, key="gen_carry_forward"):
        dept_results = st.session_state.get("dept_results", {})
        if dept_results:
            with st.spinner("🔄 " + ("正在生成承上文档..." if st.session_state.lang == "zh" else "Generating carry forward...")):
                st.session_state.carry_forward_result = generate_carry_forward(
                    all_consensus={k: v["consensus"] for k, v in dept_results.items()},
                    cross_results=st.session_state.get("cross_results", []),
                    api_url=st.session_state.api_url,
                    api_key=st.session_state.api_key,
                    model=st.session_state.model_name,
                    lang=st.session_state.lang,
                    storyboard_prompt=storyboard_text,
                    video_prompt=video_text,
                    stats=st.session_state.get("current_stats"),
                )
    
    if st.session_state.get("carry_forward_result"):
        st.success(t("carry_forward_generated"))
        st.code(st.session_state.carry_forward_result, language=None)
        
        cf_col1, cf_col2 = st.columns(2)
        with cf_col1:
            st.download_button(
                t("carry_forward_export"),
                data=st.session_state.carry_forward_result,
                file_name="carry_forward.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with cf_col2:
            st.download_button(
                t("carry_forward_copy"),
                data=st.session_state.carry_forward_result,
                file_name="carry_forward.txt",
                mime="text/plain",
                use_container_width=True,
            )
    
    st.divider()
    
    # Smart Re-roll
    with st.expander("🧠 " + ("智能回炉" if is_zh else "Smart Re-roll"), expanded=False):
        _render_smart_reroll(is_zh, is_academic=False)


def _render_smart_reroll(is_zh, is_academic=False):
    """Smart re-roll logic shared between academic and animation modes"""
    if is_academic:
        st.caption(
            "输入修改意见，AI自动分析应该回炉哪些部门，只重跑选中的部分。"
            if is_zh else
            "Enter revision feedback, AI auto-analyzes which departments to re-roll, only re-running what you select."
        )
        _placeholder = (
            "例如：方法论综述不够全面，需要补充深度学习方向；反方质疑组遗漏了重要的反对证据"
            if is_zh else
            "e.g. methodology review is incomplete, need to add deep learning direction; counter-evidence group missed key contrary evidence"
        )
    else:
        st.caption(
            "输入修改意见，AI自动分析应该回炉哪些部门。你可以在此基础上增减部门，只重跑选中的部分。"
            if is_zh else
            "Enter revision feedback, AI auto-analyzes which departments to re-roll. You can add/remove departments, only re-running what you select."
        )
        _placeholder = (
            "例如：沙虫从沙尘墙中破开冲出的瞬间不够震撼，冲击力不够；整体色调偏暖了，改成冷色沙漠感"
            if is_zh else
            "e.g. The sandworm bursting out of the dust wall isn't impactful enough; the overall tone is too warm, make it a cold desert feel"
        )
    
    revision_feedback = st.text_area(
        "✏️ " + ("修改意见" if is_zh else "Revision Feedback"),
        height=100,
        placeholder=_placeholder,
        key="smart_reroll_feedback",
    )
    
    col_analyze, _ = st.columns([2, 3])
    with col_analyze:
        if st.button("🔍 " + ("AI分析影响范围" if is_zh else "AI Analyze Impact"), type="secondary", use_container_width=True, key="analyze_reroll_impact"):
            if not revision_feedback.strip():
                st.warning("请先输入修改意见" if is_zh else "Please enter revision feedback first")
            else:
                with st.spinner("🧠 " + ("AI正在分析..." if is_zh else "AI analyzing...")):
                    impact = analyze_revision_impact(
                        revision_feedback=revision_feedback,
                        current_config=(st.session_state.get("workgroup_config") or get_current_config() or {}),
                        api_url=st.session_state.api_url,
                        api_key=st.session_state.api_key,
                        model=st.session_state.model_name,
                        lang=st.session_state.lang,
                    )
                    st.session_state.smart_reroll_impact = impact
                st.rerun()
    
    impact = st.session_state.get("smart_reroll_impact")
    if impact and not impact.get("error"):
        st.success("✅ " + ("AI分析完成！以下是推荐回炉的部门：" if is_zh else "AI analysis complete! Recommended departments to re-roll:"))
        
        available_depts = []
        _cfg = (st.session_state.get("workgroup_config") or get_current_config() or {})
        _depts = _cfg.get("departments", DEPARTMENTS)
        for dk in (_cfg.get("dept_order", list(_depts.keys())) if _cfg else DEPT_ORDER):
            dept = _depts.get(dk, {})
            name = dept.get("zh_name", dk) if is_zh else dept.get("en_name", dk)
            available_depts.append((dk, name))
        
        ai_recommended = set(d.get("dept_key") for d in impact.get("affected_departments", []))
        
        selected_depts = st.multiselect(
            "🔧 " + ("选择回炉部门" if is_zh else "Select Departments to Re-roll"),
            options=[dk for dk, _ in available_depts],
            default=[dk for dk in ai_recommended if dk in [d[0] for d in available_depts]],
            format_func=lambda dk: next((name for d, name in available_depts if d == dk), dk),
            key="reroll_dept_select",
        )
        
        if selected_depts:
            if st.button("🚀 " + ("执行回炉" if is_zh else "Execute Re-roll"), type="primary", use_container_width=True, key="exec_reroll_btn"):
                with st.spinner("🔄 " + ("正在回炉..." if is_zh else "Re-rolling...")):
                    reroll_result = smart_reroll(
                        selected_depts=selected_depts,
                        revision_feedback=revision_feedback,
                        api_url=st.session_state.api_url,
                        api_key=st.session_state.api_key,
                        model=st.session_state.model_name,
                        lang=st.session_state.lang,
                    )
                    if reroll_result:
                        st.session_state.dept_results = reroll_result.get("dept_results", st.session_state.dept_results)
                        st.session_state.final_output = reroll_result.get("final_output", st.session_state.final_output)
                        st.session_state.smart_reroll_impact = None
                        st.success("✅ " + ("智能回炉完成！请查看「最终产出」Tab" if is_zh else "Smart re-roll complete! Check Output tab"))
                        st.rerun()

def render_proofread_tab():
    is_zh = st.session_state.lang == "zh"
    
    
    
    
    # ===== Fact Check (Citation-Grounded) =====
    with st.expander("🔬 " + ("事实校验（Phase 7.5）" if is_zh else "Fact Check (Phase 7.5)"), expanded=False):
        _fc_final = st.session_state.get("final_output", {})
        _fc_report = _fc_final.get("final_report", "") or _fc_final.get("consensus_report", "") or ""
        _fc_papers = _fc_final.get("papers_data", [])

        if not _fc_report:
            st.info("需要先完成部门辩论并生成报告才能进行事实校验" if is_zh else "Complete debates and generate report first")
        else:
            # Show what we have
            _has_papers = bool(_fc_papers)
            _n_refs = len(_fc_papers)
            if _has_papers:
                st.caption(
                    f"基于{_n_refs}篇已检索文献 + 报告引用标记进行原子事实校验（RAG+NLI）"
                    if is_zh else
                    f"Atomic fact verification via RAG+NLI against {_n_refs} cached papers + citation markers"
                )
            else:
                st.caption(
                    "未找到缓存文献，将从报告参考文献部分解析并在线获取摘要"
                    if is_zh else
                    "No cached papers found, will parse bibliography and fetch abstracts online"
                )

            if st.button("🚀 " + ("开始事实校验" if is_zh else "Start Fact Check"), key="fact_check_btn"):
                # Build LLM function from session config
                api_url = st.session_state.get("api_url", "")
                api_key = st.session_state.get("api_key", "")
                model = st.session_state.get("model_name", "")

                def _llm_fn(system_prompt, user_prompt):
                    import requests as _req
                    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
                    payload = {"model": model, "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}], "temperature": 0.3}
                    try:
                        resp = _req.post(api_url, headers=headers, json=payload, timeout=120)
                        return resp.json()["choices"][0]["message"]["content"]
                    except Exception as e:
                        return f"[ERROR] {e}"

                # Build search fallback function
                def _search_fn(query, max_results=5):
                    try:
                        from academic.search_engine import AcademicSearchEngine
                        engine = AcademicSearchEngine(quality_levels=["S", "A", "B"], min_results=max_results)
                        result = engine.search(query, max_results_per_source=max(20, max_results * 4))
                        papers = result["papers"] + result.get("preprints", [])[:3]
                        return [{"title": p.title, "doi": p.doi, "abstract": p.abstract} for p in papers[:max_results]]
                    except Exception:
                        return []

                try:
                    from requirement.citation_verifier import CitationVerifier
                    _use_cv = True
                except ImportError:
                    _use_cv = False

                if _use_cv:
                    verifier = CitationVerifier(
                        llm_call_fn=_llm_fn if api_key else None,
                        search_fn=_search_fn,
                        language="zh" if is_zh else "en",
                        max_claims=20,
                        max_contexts=15,
                    )
                    with st.spinner("🔬 " + ("正在进行引用锚定原子事实校验..." if is_zh else "Running citation-grounded fact verification...")):
                        cv_report = verifier.verify(_fc_report, papers_data=_fc_papers if _fc_papers else None)
                    st.session_state.fact_check_report = cv_report
                    st.session_state.fact_check_type = "citation"
                else:
                    # Fallback to legacy FactChecker
                    checker = FactChecker(search_fn=_search_fn, llm_call_fn=_llm_fn if api_key else None)
                    # Extract claims from report
                    _paragraphs = [p.strip() for p in _fc_report.split("\n\n")
                                   if p.strip() and not p.strip().startswith("#") and len(p.strip()) >= 40]
                    _claims = [p for p in _paragraphs
                               if p.count("\n- ") <= 3 and p.count("\n* ") <= 3][:15]
                    if _claims:
                        with st.spinner("🔬 " + ("正在校验..." if is_zh else "Fact-checking...")):
                            fc_report = checker.check(_claims)
                        st.session_state.fact_check_report = fc_report
                        st.session_state.fact_check_type = "legacy"
                    else:
                        st.warning("未能提取到可验证的事实性论断" if is_zh else "No verifiable claims extracted")

            # ── Display results ──
            _fc_result = st.session_state.get("fact_check_report")
            _fc_type = st.session_state.get("fact_check_type", "legacy")

            if _fc_result and _fc_type == "citation":
                # CitationVerifier report
                st.success(f"Verification complete | Overall confidence: {_fc_result.overall_confidence:.0%}")
                st.write(_fc_result.summary)

                # Stats row
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("✅ Verified", _fc_result.verified)
                c2.metric("⚠️ Partial", _fc_result.partially_verified)
                c3.metric("❌ Contradicted", _fc_result.contradicted)
                c4.metric("❓ Unverified", _fc_result.unverified)

                # Per-claim results
                for i, cv in enumerate(_fc_result.claim_verifications):
                    status_emoji = {"verified": "✅", "partially_verified": "⚠️", "contradicted": "❌", "unverified": "❓"}
                    emoji = status_emoji.get(cv.status, "❓")
                    claim_preview = cv.claim.text[:60] + ("..." if len(cv.claim.text) > 60 else "")
                    with st.expander(f"{emoji} Claim {i+1}: {claim_preview} ({cv.status})", expanded=cv.status != "verified"):
                        st.write(f"**Claim**: {cv.claim.text}")
                        st.write(f"**Status**: {cv.status} | **Confidence**: {cv.confidence:.0%}")
                        if cv.nli_results:
                            st.write("**NLI Results**:")
                            for nli in cv.nli_results:
                                nli_emoji = {"entail": "✅", "contradict": "❌", "neutral": "➖"}.get(nli.label, "❓")
                                st.write(f"  {nli_emoji} [{nli.ref_index}] {nli.ref_title[:60]} — {nli.label} ({nli.confidence:.0%})")
                                if nli.explanation:
                                    st.caption(f"    {nli.explanation}")
                                if nli.ref_doi:
                                    st.caption(f"    DOI: {nli.ref_doi}")

            elif _fc_result and _fc_type == "legacy":
                # Legacy FactChecker report
                st.success(f"Verification complete | Overall confidence: {_fc_result.overall_confidence:.0%}")
                st.write(_fc_result.summary)
                for i, r in enumerate(_fc_result.results):
                    status_emoji = {"verified": "✅", "partially_verified": "⚠️", "contradicted": "❌", "unverified": "❓"}
                    emoji = status_emoji.get(r.status, "❓")
                    with st.expander(f"{emoji} Claim {i+1}: {r.claim[:50]}... ({r.status})", expanded=r.status != "verified"):
                        st.write(f"**Status**: {r.status} | **Confidence**: {r.confidence:.0%}")
                        if r.supporting_dois:
                            st.write("**Supporting DOI**:")
                            for doi in r.supporting_dois:
                                st.write(f"- {doi}")
                        if r.contradicting_dois:
                            st.write("**Contradicting DOI**:")
                            for doi in r.contradicting_dois:
                                st.write(f"- {doi}")
                        if r.notes:
                            st.write(f"**Notes**: {r.notes}")
    
    final = st.session_state.get("final_output", {})
    
    if not final:
        st.info(t("proofread_not_ready"))
        return
    
    # Detect pipeline mode for proofreading
    pipeline_mode = _detect_pipeline_mode()
    is_academic = pipeline_mode == "academic"
    
    st.subheader(t("proofread_result"))
    if is_academic:
        st.caption("P7: " + ("报告整合/程序/教程三部门审查" if is_zh else "Report Integration/Programming/Tutorial review"))
    else:
        st.caption("P7: " + ("空间/分镜/摄影/剪辑四部门审查" if is_zh else "Spatial/Storyboard/DP/Editing review"))
    
    # Auto-revision loop slider
    _max_loop = st.slider(
        "🔄 " + ("自动修正最大轮数" if is_zh else "Max Auto-Revision Rounds"),
        min_value=0, max_value=5, value=1,
        help=("0=关闭自动循环，1-5=校对不过时自动修正并重新校对的最大轮数" if is_zh else "0=off, 1-5=max rounds of auto revision+re-proofread loop"),
        key="auto_revision_max_loop",
    )
    
    if st.button(t("proofread_run"), type="primary", use_container_width=True, key="run_proofread"):
        dept_results = st.session_state.get("dept_results", {})
        with st.spinner(t("proofread_running")):
            if is_academic:
                report_text = final.get("final_report", "") or final.get("consensus_report", "") or ""
                result = run_academic_proofreading(
                    final_report=report_text,
                    all_consensus={k: v["consensus"] for k, v in dept_results.items()},
                    api_url=st.session_state.api_url,
                    api_key=st.session_state.api_key,
                    model=st.session_state.model_name,
                    lang=st.session_state.lang,
                    stats=st.session_state.get("current_stats"),
                )
            else:
                result = run_proofreading(
                    storyboard=final.get("storyboard_prompt", ""),
                    video_prompt=final.get("video_prompt", ""),
                    all_consensus={k: v["consensus"] for k, v in dept_results.items()},
                    api_url=st.session_state.api_url,
                    api_key=st.session_state.api_key,
                    model=st.session_state.model_name,
                    lang=st.session_state.lang,
                    stats=st.session_state.get("current_stats"),
                )
            st.session_state.proofread_result = result
            st.rerun()
    
    pr = st.session_state.get("proofread_result")
    if pr:
        st.subheader(t("proofread_overall"))
        if pr.get("passed"):
            st.success(pr.get("overall", ""))
        else:
            st.warning(pr.get("overall", ""))
        
        # Use correct department list based on pipeline mode
        _proof_departments = ACADEMIC_PROOFREAD_DEPARTMENTS if is_academic else PROOFREAD_DEPARTMENTS
        _academic_dept_names = {
            "report_integration": {"zh_name": "报告整合组", "en_name": "Report Integration"},
            "programming": {"zh_name": "程序部", "en_name": "Programming"},
            "tutorial": {"zh_name": "教程部", "en_name": "Tutorial"},
        }
        for dept_key in _proof_departments:
            if is_academic:
                dept = _academic_dept_names.get(dept_key, {})
                dept_name = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or "?"
            else:
                dept = DEPARTMENTS.get(dept_key, {})
                dept_name = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or "?"
            review = pr.get("reviews", {}).get(dept_key, "")
            with st.expander(f"🔍 **{dept_name}**"):
                st.markdown(review)
        
        # ===== Proofread Issues → Auto Fix (animation mode only) =====
        if not pr.get("passed") and not is_academic:
            st.divider()
            st.subheader(t("auto_revision"))
            
            revision = st.session_state.get("auto_revision_result")
            if revision:
                st.success(t("auto_revision_done"))
                if revision.get("revision_notes"):
                    with st.expander("📝 " + t("auto_revision_notes")):
                        st.markdown(revision["revision_notes"])
                
                # Display corrected output
                rev_sb = revision.get("storyboard_prompt", "")
                rev_vp = revision.get("video_prompt", "")
                if rev_sb:
                    with st.expander("📐 " + t("output_storyboard")):
                        st.code(rev_sb, language=None)
                if rev_vp:
                    with st.expander("🎥 " + t("output_video")):
                        st.code(rev_vp, language=None)
                
                # One-click apply correction
                if st.button("✅ " + ("应用修正结果" if is_zh else "Apply Revision"), type="primary", use_container_width=True, key="apply_auto_revision"):
                    st.session_state.final_output["storyboard_prompt"] = rev_sb
                    st.session_state.final_output["video_prompt"] = rev_vp
                    st.session_state.auto_revision_result = None
                    st.session_state.proofread_result = None
                    st.rerun()
            else:
                if st.button(t("auto_revision_run"), type="secondary", use_container_width=True, key="run_auto_revision"):
                    dept_results = st.session_state.get("dept_results", {})
                    with st.spinner(t("auto_revision_running")):
                        revision = run_auto_revision(
                            storyboard=pr.get("original_storyboard", final.get("storyboard_prompt", "")),
                            video_prompt=pr.get("original_video_prompt", final.get("video_prompt", "")),
                            proofread_result=pr,
                            all_consensus={k: v["consensus"] for k, v in dept_results.items()},
                            api_url=st.session_state.api_url,
                            api_key=st.session_state.api_key,
                            model=st.session_state.model_name,
                            lang=st.session_state.lang,
                            stats=st.session_state.get("current_stats"),
                        )
                        st.session_state.auto_revision_result = revision
                        st.rerun()

        # ===== Academic Auto Revision (loop) =====
        # Show loop results whenever they exist (even if passed=True after loop)
        _ac_loop_result = st.session_state.get("academic_loop_result")
        if _ac_loop_result:
            st.divider()
            rounds_done = _ac_loop_result.get("rounds_done", 0)
            final_passed = _ac_loop_result.get("passed", False)
            if final_passed:
                st.success(f"✅ {('自动修正通过！共' if is_zh else 'Auto-revision passed! ')}{rounds_done}{('轮' if is_zh else ' rounds')}")
            else:
                st.warning(f"⚠️ {('达到最大轮数限制（' if is_zh else 'Max rounds reached (')}{_max_loop}{('轮），校对仍有问题' if is_zh else '), issues remain')}")
            
            # Show each round summary
            for ri, rinfo in enumerate(_ac_loop_result.get("rounds", []), 1):
                with st.expander(f"🔄 {('第' if is_zh else 'Round ')}{ri}{('轮' if is_zh else '')} — {'✅' if rinfo.get('passed') else '❌'}"):
                    if rinfo.get("revision_notes"):
                        st.markdown(f"**{('修正说明' if is_zh else 'Revision Notes')}:**\n{rinfo['revision_notes']}")
                    if rinfo.get("issues_summary"):
                        st.markdown(f"**{('校对结果' if is_zh else 'Proofread Result')}:**\n{rinfo['issues_summary']}")
            
            # Show final revised report + apply button
            rev_report = _ac_loop_result.get("final_report", "")
            if rev_report:
                with st.expander("📄 " + ("修正后报告" if is_zh else "Revised Report")):
                    st.markdown(rev_report)
            if st.button("✅ " + ("应用修正结果" if is_zh else "Apply Revision"), type="primary", use_container_width=True, key="apply_ac_loop_revision"):
                st.session_state.final_output["final_report"] = rev_report
                st.session_state.academic_loop_result = None
                st.session_state.proofread_result = None
                st.rerun()
        elif not pr.get("passed") and is_academic:
            st.divider()
            if _max_loop >= 1:
                if st.button("🔄 " + (f"开始自动修正循环（最多{_max_loop}轮）" if is_zh else f"Start Auto-Revision Loop (max {_max_loop} rounds)"), type="secondary", use_container_width=True, key="run_ac_loop"):
                        dept_results = st.session_state.get("dept_results", {})
                        report_text = final.get("final_report", "") or final.get("consensus_report", "") or ""
                        current_report = report_text
                        current_pr = pr
                        rounds_log = []
                        
                        with st.status("🔄 " + ("自动修正循环中..." if is_zh else "Auto-revision loop running..."), expanded=True) as status:
                            for round_i in range(1, _max_loop + 1):
                                st.write(f"--- **{('第' if is_zh else 'Round ')}{round_i}{('轮' if is_zh else '')}** ---")
                                
                                # Step 1: Auto-revise
                                st.write("🔧 " + ("修正中..." if is_zh else "Revising..."))
                                ac_revision = run_academic_auto_revision(
                                    final_report=current_report,
                                    proofread_result=current_pr,
                                    all_consensus={k: v["consensus"] for k, v in dept_results.items()},
                                    api_url=st.session_state.api_url,
                                    api_key=st.session_state.api_key,
                                    model=st.session_state.model_name,
                                    lang=st.session_state.lang,
                                    stats=st.session_state.get("current_stats"),
                                )
                                current_report = ac_revision.get("final_report", current_report)
                                
                                # Step 2: Re-proofread
                                st.write("🔍 " + ("重新校对中..." if is_zh else "Re-proofreading..."))
                                new_pr = run_academic_proofreading(
                                    final_report=current_report,
                                    all_consensus={k: v["consensus"] for k, v in dept_results.items()},
                                    api_url=st.session_state.api_url,
                                    api_key=st.session_state.api_key,
                                    model=st.session_state.model_name,
                                    lang=st.session_state.lang,
                                    stats=st.session_state.get("current_stats"),
                                )
                                current_pr = new_pr
                                
                                # Log this round
                                rounds_log.append({
                                    "passed": new_pr.get("passed", False),
                                    "revision_notes": ac_revision.get("revision_notes", ""),
                                    "issues_summary": new_pr.get("overall", ""),
                                })
                                
                                if new_pr.get("passed"):
                                    st.write("✅ " + ("校对通过！" if is_zh else "Proofreading passed!"))
                                    status.update(label="✅ " + ("自动修正完成" if is_zh else "Auto-revision complete"), state="complete")
                                    break
                                else:
                                    remaining = _max_loop - round_i
                                    st.write("❌ " + (f"仍有问题，剩余{remaining}轮" if is_zh else f"Issues remain, {remaining} rounds left"))
                            else:
                                status.update(label="⚠️ " + ("达到最大轮数" if is_zh else "Max rounds reached"), state="complete")
                        
                        st.session_state.academic_loop_result = {
                            "rounds_done": len(rounds_log),
                            "passed": current_pr.get("passed", False),
                            "rounds": rounds_log,
                            "final_report": current_report,
                        }
                        st.session_state.proofread_result = current_pr
                        st.rerun()
            else:
                st.info("ℹ️ " + ("滑块设为0，不会自动修正。请手动点击上方修正按钮或调整滑块。" if is_zh else "Slider set to 0, auto-revision disabled. Use manual revision or adjust slider."))
        
        # ===== Director Instruction Correction =====
        st.divider()
        st.subheader(t("director_revision"))
        st.caption("导演指定部门+修改意见 → 重跑该部门辩论 → 重新生成分镜表+视频提示词" if is_zh else "Director picks department + notes → re-debate that dept → regenerate storyboard + video prompt")
        
        dept_options = {}
        for dk in DEPT_ORDER:
            dept = DEPARTMENTS.get(dk, {})
            dept_options[dk] = (dept.get("zh_name") if is_zh else dept.get("en_name")) or dept.get("zh_name") or dept.get("en_name") or "?"
        
        col_dept, col_rounds = st.columns([3, 1])
        with col_dept:
            selected_dept = st.selectbox(
                t("director_revision_dept"),
                options=list(dept_options.keys()),
                format_func=lambda x: dept_options[x],
                key="director_revision_dept_select",
            )
        with col_rounds:
            revision_rounds = st.number_input(
                t("revision_rounds"),
                min_value=1, max_value=5, value=2,
                key="director_revision_rounds",
            )
        
        director_note = st.text_area(
            t("director_revision_note"),
            height=100,
            placeholder=t("director_revision_note_hint"),
            key="director_revision_note_input",
        )
        
        if st.button(t("director_revision_run"), type="primary", use_container_width=True, key="run_director_revision"):
            if not director_note.strip():
                st.warning(t("reject_hint"))
            else:
                dept_results = st.session_state.get("dept_results", {})
                with st.spinner(t("director_revision_running", dept=dept_options[selected_dept])):
                    result = run_director_revision(
                        department_key=selected_dept,
                        director_note=director_note,
                        all_consensus={k: v["consensus"] for k, v in dept_results.items()},
                        user_script=st.session_state.script,
                        positive_prompt=st.session_state.positive_prompt,
                        negative_prompt=st.session_state.negative_prompt,
                        character_refs=st.session_state.character_refs,
                        cross_results=st.session_state.get("cross_results", []),
                        current_storyboard=final.get("storyboard_prompt", ""),
                        current_video_prompt=final.get("video_prompt", ""),
                        api_url=st.session_state.api_url,
                        api_key=st.session_state.api_key,
                        model=st.session_state.model_name,
                        lang=st.session_state.lang,
                        rounds=revision_rounds,
                        stats=st.session_state.get("current_stats"),
                    )
                    # Update department consensus and final output
                    st.session_state.dept_results[selected_dept] = result["dept_result"]
                    st.session_state.final_output = {
                        "storyboard_prompt": result["storyboard_prompt"],
                        "video_prompt": result["video_prompt"],
                    }
                    # Clear old proofread results
                    st.session_state.proofread_result = None
                    st.session_state.auto_revision_result = None
                    st.rerun()



# ============ Comparison Panel ============

def render_compare_tab():
    """Render run comparison panel"""
    is_zh = st.session_state.lang == "zh"
    
    
    
    
    run_history = st.session_state.get("run_history", [])
    
    st.subheader(t("compare_title"))
    st.caption(t("compare_desc"))
    
    # Save current run results
    if st.session_state.get("debate_completed") and st.session_state.get("final_output"):
        if st.button(t("compare_save_run"), type="primary", use_container_width=True, key="save_current_run"):
            stats = st.session_state.get("current_stats", {})
            start_time = st.session_state.get("current_start_time")
            end_time = st.session_state.get("current_end_time", time.time())
            duration = round(end_time - start_time, 1) if start_time else 0
            
            run_record = {
                "timestamp": datetime.now().isoformat(),
                "model_profile": st.session_state.get("model_profile", "custom"),
                "model_name": st.session_state.model_name,
                "architecture_mode": st.session_state.architecture_mode,
                "mode_name": ARCHITECTURE_MODES.get(st.session_state.architecture_mode, {}).get("zh_name" if is_zh else "en_name", ""),
                "prompt_tokens": stats.get("prompt_tokens", 0),
                "completion_tokens": stats.get("completion_tokens", 0),
                "total_tokens": stats.get("total_tokens", 0),
                "api_calls": stats.get("api_calls", 0),
                "duration_seconds": duration,
                "storyboard_prompt": st.session_state.final_output.get("storyboard_prompt", ""),
                "video_prompt": st.session_state.final_output.get("video_prompt", ""),
                "script_hash": hash(st.session_state.script) if st.session_state.script else 0,
            }
            run_history.append(run_record)
            st.session_state.run_history = run_history
            st.success(t("compare_saved"))
    
    if not run_history:
        st.info(t("compare_no_runs"))
        return
    
    # Comparison overview table
    st.subheader(t("compare_table"))
    
    # Build comparison data
    cols = ["#", t("compare_model"), t("compare_mode"), t("compare_tokens"), t("compare_time"), t("compare_api_calls")]
    table_data = []
    for i, run in enumerate(run_history):
        mode_name = ARCHITECTURE_MODES.get(run["architecture_mode"], {}).get("zh_name" if is_zh else "en_name", run["architecture_mode"])
        table_data.append({
            "#": i + 1,
            t("compare_model"): run["model_name"],
            t("compare_mode"): mode_name,
            t("compare_tokens"): f"{run['total_tokens']:,}",
            t("compare_time"): f"{run['duration_seconds']}s",
            t("compare_api_calls"): run["api_calls"],
        })
    
    st.dataframe(table_data, use_container_width=True, hide_index=True)
    
    # Details for each run
    for i, run in enumerate(run_history):
        mode_name = ARCHITECTURE_MODES.get(run["architecture_mode"], {}).get("zh_name" if is_zh else "en_name", run["architecture_mode"])
        with st.expander(f"\u2615 {t('compare_run_label', idx=i+1)} \u2014 {run['model_name']} / {mode_name}"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric(t("compare_tokens"), f"{run['total_tokens']:,}")
            with col2:
                st.metric(t("compare_time"), f"{run['duration_seconds']}s")
            with col3:
                st.metric(t("compare_api_calls"), run["api_calls"])
            
            st.divider()
            
            sb = run.get("storyboard_prompt", "")
            vp = run.get("video_prompt", "")
            if sb:
                with st.expander("\U0001F4D0 " + t("compare_storyboard_preview")):
                    st.code(sb[:2000] + ("..." if len(sb) > 2000 else ""), language=None)
            if vp:
                with st.expander("\U0001F3A5 " + t("compare_video_preview")):
                    st.code(vp[:2000] + ("..." if len(vp) > 2000 else ""), language=None)
    
    # Clear button
    col_clear1, col_clear2 = st.columns([3, 1])
    with col_clear2:
        if st.button(t("compare_clear"), key="clear_all_runs"):
            st.session_state.run_history = []
            st.rerun()


# ============ Main Entry ============


# ============ Market Tab ============

def render_market_tab():
    """Render Market Tab"""
    is_zh = st.session_state.lang == "zh"
    
    
    
    
    st.subheader(t("market_title"))
    st.caption(t("market_subtitle"))
    
    if not st.session_state.api_key or not st.session_state.script.strip():
        st.warning(t("market_need_debate"))
        return
    
    # View existing results
    market_result = st.session_state.get("market_result")
    
    # 🔧 Auto-recover: restore from disk after network loss/crash
    if not market_result:
        # Check for disk archives
        saved = autosave_load("market_result")
        if saved:
            st.info("💾 " + ("检测到上次未完成的市场模式结果，已自动恢复！" if is_zh else "Previous market result recovered from disk!"))
            st.session_state.market_result = saved
            market_result = saved
        else:
            # Check for staged archives
            step3 = autosave_load("market_step3")
            step2 = autosave_load("market_step2")
            step1 = autosave_load("market_step1")
            if step3:
                st.warning("⚠️ " + ("检测到投票完成但补丁未完成的结果，已恢复到投票阶段" if is_zh else "Vote completed but patch not done, recovered to vote stage"))
                st.session_state.market_result = step3
                market_result = step3
            elif step2:
                st.warning("⚠️ " + ("检测到候选+出题完成但投票未完成的结果，已恢复" if is_zh else "Candidates+questions done but vote not completed, recovered"))
                st.session_state.market_result = step2
                market_result = step2
            elif step1:
                st.warning("⚠️ " + ("检测到候选生成完成但出题未完成的结果，已恢复" if is_zh else "Candidates generated but questions not done, recovered"))
                st.session_state.market_result = step1
                market_result = step1
    
    if not market_result:
        # Import existing result option
        existing_result = st.session_state.get("final_output")
        if existing_result and existing_result.get("storyboard_prompt"):
            include_existing = st.checkbox(
                "📥 " + ("将普通模式已有结果作为候选之一参加竞选" if is_zh else "Include normal mode result as a candidate"),
                value=True,
                key="market_include_existing",
            )
            if include_existing:
                st.caption("✅ " + ("普通模式结果将作为候选A参与市场竞选，其余候选从API生成" if is_zh else "Normal mode result will join as Candidate A, other candidates generated via API"))
        
        # Department filter — fewer departments = less memory
        # Use current config (not module-level DEPT_ORDER) so academic mode shows correct departments
        _mkt_cfg = st.session_state.get("workgroup_config") or get_current_config() or {}
        _mkt_cfg_depts = _mkt_cfg.get("departments", {})
        _mkt_cfg_order = _mkt_cfg.get("dept_order", [])
        _all_dept_keys = _mkt_cfg_order if _mkt_cfg_order else (list(_mkt_cfg_depts.keys()) if _mkt_cfg_depts else list(DEPARTMENTS.keys()))
        _dept_labels = {k: _mkt_cfg_depts.get(k, DEPARTMENTS.get(k, {})).get("zh_name" if is_zh else "en_name", k) for k in _all_dept_keys}
        # Clear stale widget cache when config changes (e.g. animation → academic)
        _cfg_fp = str(sorted(_all_dept_keys))
        if st.session_state.get("_market_dept_cfg_fp") != _cfg_fp:
            st.session_state.pop("market_dept_filter", None)
            st.session_state._market_dept_cfg_fp = _cfg_fp
        _selected_depts = st.multiselect(
            "🏷️ " + ("选择要跑的部门（越少越省内存）" if is_zh else "Select departments (fewer = less memory)"),
            options=_all_dept_keys,
            default=_all_dept_keys,
            format_func=lambda k: _dept_labels.get(k, k),
            key="market_dept_filter",
        )
        if not _selected_depts:
            st.warning("⚠️ " + ("至少选择一个部门" if is_zh else "Select at least one department"))

        # Start button
        if st.button(t("market_start"), type="primary", use_container_width=True):
            run_market()
            st.rerun()
        return
    
    # ---- Display results ----
    
    # Step 1: Candidate proposals
    candidates = market_result.get("candidates", [])
    st.subheader(t("market_candidates"))
    for c in candidates:
        existing_tag = " ⭐" + ("已有" if is_zh else "Existing") if c.get("is_existing") else f" (T={c.get('temperature', '?')})"
        with st.expander(f"📦 {t('market_candidates')} {c.get('label', '?')}{existing_tag}"):
            st.markdown(f"**分镜表**" if is_zh else "**Storyboard**")
            st.code(c.get("storyboard_prompt", "")[:500] + "...", language=None)
            st.markdown(f"**视频提示词**" if is_zh else "**Video Prompt**")
            st.code(c.get("video_prompt", "")[:500] + "...", language=None)
    
    st.divider()
    
    # Step 2: Issue statistics (may not exist)
    questions = market_result.get("questions", [])
    if questions:
        st.subheader(t("market_questions_total", total=len(questions)))
        
        # Count issues by department
        dept_q_counts = {}
        for q in questions:
            dk = q.get("source_dept", "unknown")
            dn = DEPARTMENTS.get(dk, {}).get("zh_name" if is_zh else "en_name", dk)
            dept_q_counts[dn] = dept_q_counts.get(dn, 0) + 1
        
        if dept_q_counts:
            st.bar_chart(dept_q_counts)
        
        st.divider()
    else:
        st.info("❓ " + ("出题尚未完成，重新开市将从此步骤继续" if is_zh else "Questions not yet generated, restart market to continue from this step"))
        # Offer re-open market button
        if st.button("🔄 " + ("重新开市（从出题开始）" if is_zh else "Restart market (from questions)"), type="primary"):
            # Keep candidates, re-run question generation
            st.session_state.market_result = None
            # Candidate data preserved in step1 archive
            st.rerun()
        return
    
    # Step 3: Vote results (may not exist)
    vote_result = market_result.get("vote_result")
    if vote_result:
        vote_counts = vote_result.get("vote_counts", {})
        winner_label = vote_result.get("winner_label", "?")
        winner_votes = vote_counts.get(winner_label, 0)
        
        st.subheader(t("market_vote_result"))
        st.success(t("market_winner", label=winner_label, votes=winner_votes))
        
        # Vote count bar chart
        if vote_counts:
            st.bar_chart(vote_counts)
        
        # Vote details
        with st.expander("📊 " + t("market_vote_detail")):
            votes = vote_result.get("votes", [])
            # Display grouped by department
            dept_votes = {}
            for v in votes:
                voter_dept = v.get("voter", "").split("/")[0] if "/" in v.get("voter", "") else "Other"
                if voter_dept not in dept_votes:
                    dept_votes[voter_dept] = []
                dept_votes[voter_dept].append(v)
            
            for dept_name, dept_v in dept_votes.items():
                with st.expander(f"🏛️ {dept_name}"):
                    for v in dept_v[:10]:  # 每部门最多显示10个
                        st.markdown(f"- **Q:** {v['question'][:80]}... → **Choice:** Candidate {v.get('choice', '?')} {v.get('reason', '')[:100]}")
        
        st.divider()
        
        # Controversial issues
        contested = vote_result.get("contested_questions", [])
        if contested:
            with st.expander("⚠️ " + t("market_contested"), expanded=True):
                st.caption(t("market_contested_hint"))
                for cq in contested[:10]:
                    st.markdown(f"- {cq['question']}（{cq['vote_distribution']}）")
        
        st.divider()
    else:
        st.info("🗳️ " + ("投票尚未完成，重新开市将从此步骤继续" if is_zh else "Voting not yet completed, restart market to continue from this step"))
        if st.button("🔄 " + ("重新开市（从投票开始）" if is_zh else "Restart market (from voting)"), type="primary"):
            st.session_state.market_result = None
            st.rerun()
        return
    
    # Step 4: Patch correction
    patch_result = market_result.get("patch_result")
    if patch_result:
        st.subheader(t("market_patch_result"))
        st.info(patch_result.get("patch_notes", ""))
        
        with st.expander("📐 " + ("补丁版分镜表" if is_zh else "Patched Storyboard")):
            st.code(patch_result.get("storyboard_prompt", ""), language=None)
        with st.expander("🎥 " + ("补丁版视频提示词" if is_zh else "Patched Video Prompt")):
            st.code(patch_result.get("video_prompt", ""), language=None)
    
    st.divider()
    
    # Apply button
    col1, col2 = st.columns(2)
    with col1:
        if st.button(t("market_apply_winner"), type="primary", use_container_width=True):
            winner = vote_result.get("winner_candidate", {})
            st.session_state.final_output = {
                "storyboard_prompt": winner.get("storyboard_prompt", ""),
                "video_prompt": winner.get("video_prompt", ""),
            }
            st.session_state.debate_completed = True
            st.success(t("market_applied"))
    
    with col2:
        if patch_result and st.button(t("market_apply_patched"), use_container_width=True):
            st.session_state.final_output = {
                "storyboard_prompt": patch_result.get("storyboard_prompt", ""),
                "video_prompt": patch_result.get("video_prompt", ""),
            }
            st.session_state.debate_completed = True
            st.success(t("market_applied"))


def run_market():
    """Execute full market flow — fully protected, no step failure wastes previous results"""
    is_zh = st.session_state.lang == "zh"
    api_url = st.session_state.api_url
    api_key = st.session_state.api_key
    model = st.session_state.model_name
    lang = st.session_state.lang
    market_rounds = st.session_state.get("market_rounds", 3)
    extra = st.session_state.extra_instructions
    num_candidates = st.session_state.market_num_candidates
    questions_per = st.session_state.market_questions_per
    
    stats = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "api_calls": 0}
    st.session_state.current_stats = stats
    st.session_state.current_start_time = time.time()
    
    progress = st.progress(0, text=t("market_running"))
    
    # ====== Step 1: Generate candidates ======
    progress.progress(0.05, text=t("market_step1"))
    
    completed_count = 0
    def on_gen_progress(phase, idx, dept_key, di=0, total_depts=8):
        nonlocal completed_count
        try:
            label = chr(65 + idx)
            if phase == "candidate_done":
                completed_count += 1
                progress.progress(min(0.1 + 0.4 * (completed_count / num_candidates), 0.49), 
                    text=f"✅ {('候选' if is_zh else 'Cand.')} {label} {('已完成' if is_zh else 'done')}（{completed_count}/{num_candidates}）")
            elif phase == "candidate_error":
                progress.progress(min(0.1 + 0.4 * (completed_count / num_candidates), 0.49),
                    text=f"❌ {('候选' if is_zh else 'Cand.')} {label}: {dept_key[:60]}")
            elif phase == "dept_done":
                dept = DEPARTMENTS.get(dept_key, {})
                dn = dept.get("zh_name" if is_zh else "en_name", dept_key)
                pct_val = min(0.1 + 0.4 * (completed_count / num_candidates) + (di / total_depts) / num_candidates, 0.49)
                progress.progress(pct_val, text=f"📦 {('候选' if is_zh else 'Cand.')} {label}: {dn} ✅（{di}/{total_depts}）")
            elif phase == "dept_start":
                dept = DEPARTMENTS.get(dept_key, {})
                dn = dept.get("zh_name" if is_zh else "en_name", dept_key)
                pct_val = min(0.1 + 0.4 * (completed_count / num_candidates) + (di / total_depts) / num_candidates, 0.49)
                progress.progress(pct_val, text=f"📦 {('候选' if is_zh else 'Cand.')} {label}: {dn} {('辩论中' if is_zh else 'debating')}（{di+1}/{total_depts}）")
            elif phase == "phase":
                phase_names = {
                    "spatial_review": ("空间审核" if is_zh else "Spatial Review"),
                    "cross_debate": ("交叉辩论" if is_zh else "Cross Debate"),
                    "summary": ("总结" if is_zh else "Summary"),
                }
                pn = phase_names.get(dept_key, dept_key)
                pct_val = min(0.1 + 0.4 * (completed_count / num_candidates) + 0.35 / num_candidates, 0.49)
                progress.progress(pct_val, text=f"📦 {('候选' if is_zh else 'Cand.')} {label}: {pn}...")
        except Exception:
            pass  # 进度更新失败不影响主流程
    
    # If importing existing result, API generates one fewer candidate
    include_existing = st.session_state.get("market_include_existing", False)
    existing_result = st.session_state.get("final_output")
    has_existing = include_existing and existing_result and existing_result.get("storyboard_prompt")
    api_candidates = max(num_candidates - 1, 1) if has_existing else num_candidates

    # Department filter from market UI — validate against current config
    _dept_filter = st.session_state.get("market_dept_filter")
    if _dept_filter:
        _mkt_cfg = st.session_state.get("workgroup_config") or get_current_config() or {}
        _valid_depts = set(_mkt_cfg.get("dept_order", [])) or set(_mkt_cfg.get("departments", {}).keys())
        _dept_filter = [d for d in _dept_filter if d in _valid_depts]
        if not _dept_filter:
            _dept_filter = None  # fallback to all if nothing valid
    else:
        _dept_filter = None  # None = all departments

    try:
        candidates_result = generate_candidates(
            num_candidates=api_candidates,
            user_script=st.session_state.script,
            positive_prompt=st.session_state.positive_prompt,
            negative_prompt=st.session_state.negative_prompt,
            character_refs=st.session_state.character_refs,
            api_url=api_url, api_key=api_key, model=model,
            rounds=market_rounds, lang=lang, extra_instructions=extra,
            carry_forward=st.session_state.carry_forward,
            stats=stats,
            progress_callback=on_gen_progress,
            max_workers=st.session_state.market_parallel_workers,
            dept_filter=_dept_filter,
        )
    except Exception as e:
        st.error(f"❌ Step 1 {('生成候选失败' if is_zh else 'Candidate generation failed')}: {e}")
        progress.progress(1.0, text="❌ " + ("失败" if is_zh else "Failed"))
        return
    
    candidates = candidates_result.get("candidates", [])
    if not candidates:
        st.error("❌ " + ("未生成任何候选方案，请检查API配置" if is_zh else "No candidates generated, check API config"))
        progress.progress(1.0, text="❌ " + ("失败" if is_zh else "Failed"))
        return
    
    # Import normal mode existing result as candidate
    include_existing = st.session_state.get("market_include_existing", False)
    existing_result = st.session_state.get("final_output")
    if include_existing and existing_result and existing_result.get("storyboard_prompt"):
        existing_candidate = {
            "label": "A",
            "temperature": 0.0,
            "storyboard_prompt": existing_result.get("storyboard_prompt", ""),
            "video_prompt": existing_result.get("video_prompt", ""),
            "dept_results": existing_result.get("dept_results", {}),
            "stats": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "api_calls": 0},
            "is_existing": True,
        }
        # Existing result as candidate A, remaining candidates start from B
        candidates.insert(0, existing_candidate)
        for i, c in enumerate(candidates[1:], 1):
            c["label"] = chr(65 + i)  # B, C, D...
        progress.progress(0.49, text="📥 " + ("已导入普通模式结果作为候选A" if is_zh else "Imported normal mode result as Candidate A"))
    
    step1_errors = candidates_result.get("errors", [])
    if step1_errors:
        st.warning(f"⚠️ {len(step1_errors)} {('个候选生成失败' if is_zh else 'candidates failed')}")
    
    # 📦 Step1 complete → save to disk
    autosave_result("market_step1", {"candidates": candidates, "errors": step1_errors})
    # ====== Step 2: Question generation ======
    progress.progress(0.55, text=t("market_step2"))
    
    def on_q_progress(phase, dept_key, debater_key):
        try:
            dept = DEPARTMENTS.get(dept_key, {})
            dn = dept.get("zh_name" if is_zh else "en_name", dept_key)
            # debater_key now passes dept_key (dept question mode), use dept name directly
            progress.progress(0.55, text=f"❓ {dn} {('出题中' if is_zh else 'questioning')}...")
        except Exception:
            pass
    
    try:
        questions_result = generate_questions(
            candidates=candidates,
            questions_per_debater=questions_per,
            api_url=api_url, api_key=api_key, model=model,
            lang=lang, stats=stats,
            progress_callback=on_q_progress,
        )
    except Exception as e:
        st.error(f"❌ Step 2 {('出题失败' if is_zh else 'Question generation failed')}: {e}")
        # Keep candidate results even if question generation fails
        st.session_state.market_result = {
            "candidates": candidates, "questions": [],
            "vote_result": None, "patch_result": None,
        }
        st.session_state.current_end_time = time.time()
        autosave_result("market_result", st.session_state.market_result)
        progress.progress(1.0, text="⚠️ " + ("候选已保存，但出题失败" if is_zh else "Candidates saved, questioning failed"))
        return
    
    questions = questions_result.get("questions", [])
    if not questions:
        st.warning("⚠️ " + ("未生成评估问题，跳过投票" if is_zh else "No questions generated, skipping vote"))
        st.session_state.market_result = {
            "candidates": candidates, "questions": [],
            "vote_result": None, "patch_result": None,
        }
        st.session_state.current_end_time = time.time()
        autosave_result("market_result", st.session_state.market_result)
        progress.progress(1.0, text="⚠️ " + ("候选已保存" if is_zh else "Candidates saved"))
        return
    
    # 📦 Step2 complete → save to disk
    autosave_result("market_step2", {"candidates": candidates, "questions": questions})
    
    # ====== Step 3: Voting ======
    progress.progress(0.7, text=t("market_step3"))
    
    def on_vote_progress(phase, vi, debater_name, total_voters=None):
        try:
            tv = total_voters if total_voters else 24
            pct = min(0.7 + 0.2 * (vi / max(tv, 1)), 0.89)
            progress.progress(pct, text=t("market_voting", voter=debater_name))
        except Exception:
            pass
    
    try:
        vote_result = vote_on_questions(
            questions=questions,
            candidates=candidates,
            api_url=api_url, api_key=api_key, model=model,
            lang=lang, stats=stats,
            progress_callback=on_vote_progress,
        )
    except Exception as e:
        st.error(f"❌ Step 3 {('投票失败' if is_zh else 'Voting failed')}: {e}")
        st.session_state.market_result = {
            "candidates": candidates, "questions": questions,
            "vote_result": None, "patch_result": None,
        }
        st.session_state.current_end_time = time.time()
        autosave_result("market_result", st.session_state.market_result)
        progress.progress(1.0, text="⚠️ " + ("候选已保存，投票失败" if is_zh else "Candidates saved, voting failed"))
        return
    
    # Check voting results
    winner = vote_result.get("winner_candidate") if vote_result else None
    if not winner and candidates:
        # Vote failed but candidates exist, take first one
        winner = candidates[0]
        st.warning("⚠️ " + ("投票未产生明确胜者，默认使用候选A" if is_zh else "No clear winner, using Candidate A"))
    
    # 📦 Step3 complete → save to disk
    autosave_result("market_step3", {"candidates": candidates, "questions": questions, "vote_result": vote_result})
    
    # ====== Step 4: Patch correction ======
    progress.progress(0.92, text=t("market_step4"))
    
    patch_result = None
    if winner and vote_result and vote_result.get("contested_questions"):
        try:
            patch_result = patch_winner(
                winner_candidate=winner,
                contested_questions=vote_result["contested_questions"],
                all_candidates=candidates,
                api_url=api_url, api_key=api_key, model=model,
                lang=lang, stats=stats,
            )
        except Exception as e:
            st.warning(f"⚠️ {('补丁修正失败' if is_zh else 'Patch failed')}: {e}")
            patch_result = None
    elif winner:
        patch_result = {
            "storyboard_prompt": winner.get("storyboard_prompt", ""),
            "video_prompt": winner.get("video_prompt", ""),
            "patch_notes": ("无争议问题，无需修正" if is_zh else "No contested issues"),
        }
    
    # ====== Complete ======
    progress.progress(1.0, text="✅ " + ("市场收市！" if is_zh else "Market closed!"))
    
    st.session_state.market_result = {
        "candidates": candidates,
        "questions": questions,
        "vote_result": vote_result,
        "patch_result": patch_result,
    }
    
    st.session_state.current_end_time = time.time()
    # 📦 Final result saved to disk
    autosave_result("market_result", st.session_state.market_result)


# ============ v3.0 Smart Grouping Tab ============

def render_config_tab():
    """🧠 Smart Grouping — Consensus Pipeline v3.0 Entry"""
    is_zh = st.session_state.lang == "zh"

    
    

    # Ensure workgroup_config is initialized
    if not st.session_state.get("workgroup_config"):
        st.session_state.workgroup_config = get_current_config() or {}
    
    # Detect academic vs animation mode for conditional UI
    _cfg_depts = (st.session_state.get("workgroup_config") or get_current_config() or {}).get("departments", {})
    is_academic = _detect_pipeline_mode() == "academic"
    is_animation = not is_academic
    
    st.subheader(t("config_title"))
    st.caption(t("config_subtitle"))
    
    # Display current config
    _cur_name = get_current_config_name()
    # Use workgroup_config from session_state when available (more accurate than global DEPARTMENTS)
    _display_depts = (st.session_state.get("workgroup_config") or {}).get("departments", DEPARTMENTS)
    _total_debaters = sum(len(d.get("debaters", {})) if isinstance(d, dict) else 0 for d in _display_depts.values()) if _display_depts else 0
    st.info(f"{'当前配置' if is_zh else 'Current config'}: **{_cur_name}**  |  {'部门数' if is_zh else 'Depts'}: {len(_display_depts)}  |  {'辩手总数' if is_zh else 'Total debaters'}: {_total_debaters}")
    
    left_col, right_col = st.columns([2, 3])
    
    with left_col:
        st.markdown("### " + ("需求输入" if is_zh else "Requirement Input"))
        
        # User requirement description
        # Allow using input tab content as config input
        _script_content = st.session_state.get("script", "")
        if _script_content:
            _btn_label = "📋 " + ("使用输入Tab内容" if is_zh else "Use Input Tab Content")
            if st.button(_btn_label, use_container_width=True, key="btn_use_input_tab"):
                st.session_state.config_user_input = _script_content
                st.rerun()

        user_input = st.text_area(
            t("config_input_label"),
            height=150,
            placeholder=t("config_input_hint"),
            key="config_user_input",
        )
        
        # AI smart grouping button
        if st.button(t("config_ai_button"), type="primary", use_container_width=True, disabled=not user_input):
            if not st.session_state.api_key:
                st.error(t("need_api"))
            else:
                with st.spinner(t("config_ai_running")):
                    result = analyze_and_configure(
                        user_input=user_input,
                        api_url=st.session_state.api_url,
                        api_key=st.session_state.api_key,
                        model=st.session_state.model_name,
                        lang=st.session_state.lang,
                    )
                if result.get("error"):
                    st.error(f"{t('config_error')}: {result.get('message', '')}")
                    if result.get("raw_content"):
                        with st.expander("原始返回" if is_zh else "Raw Response"):
                            st.code(result["raw_content"])
                elif result.get("needs_clarification"):
                    st.warning(t("config_clarification"))
                    for q in result.get("clarification_questions", []):
                        st.markdown(f"- {q}")
                    # Still save partial config for editing
                    apply_config(result)
                    st.session_state.workgroup_config = result
                else:
                    apply_config(result)
                    st.session_state.workgroup_config = result
                    st.success("✅ " + ("AI配组完成！" if is_zh else "AI config complete!"))
        
        st.divider()
        
        # Preset selection
        st.markdown("### " + t("config_preset_label"))
        presets = list_presets()
        if presets:
            selected_preset = st.selectbox(
                t("config_preset_label"),
                options=presets,
                key="config_preset_select",
            )
            if st.button("📦 " + ("加载预设" if is_zh else "Load Preset"), use_container_width=True):
                try:
                    cfg = load_preset(selected_preset)
                    apply_config(cfg)
                    st.session_state.workgroup_config = cfg
                    st.session_state.workgroup_name = selected_preset
                    st.session_state.pipeline_mode = None  # reset, will be detected on next use
                    st.success(f"✅ {'已加载预设' if is_zh else 'Loaded preset'}: {selected_preset}")
                except Exception as e:
                    st.error(str(e))
        
        # User config selection
        st.markdown("### " + t("config_profile_label"))
        profiles = list_profiles()
        if profiles:
            selected_profile = st.selectbox(
                t("config_profile_label"),
                options=profiles,
                key="config_profile_select",
            )
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button("💾 " + ("加载" if is_zh else "Load"), use_container_width=True):
                    try:
                        cfg = load_profile(selected_profile)
                        apply_config(cfg)
                        st.session_state.workgroup_config = cfg
                        st.session_state.workgroup_name = selected_profile
                        st.success(f"✅ {'已加载配置' if is_zh else 'Loaded config'}: {selected_profile}")
                    except Exception as e:
                        st.error(str(e))
            with col_b:
                if st.button(t("config_delete"), use_container_width=True):
                    delete_profile(selected_profile)
                    st.rerun()
    
    with right_col:
        st.markdown("### " + t("config_preview_title"))
        
        config = st.session_state.get("workgroup_config")
        if config is None:
            # Try loading currently active config
            config = get_current_config()
        
        if not config or not config.get("departments"):
            st.info(t("config_no_depts"))
        else:
            # One collapsible section per department
            depts = config.get("departments", {})
            dept_order = config.get("dept_order", list(depts.keys()))
            
            for dept_key in dept_order:
                if dept_key not in depts:
                    continue
                dept = depts[dept_key]
                dept_name = dept.get("zh_name", dept_key) if is_zh else dept.get("en_name", dept_key)
                debaters = dept.get("debaters", {})
                if not isinstance(debaters, dict):
                    debaters = {}
                num_debaters = len(debaters)
                
                with st.expander(f"**{dept_name}** ({num_debaters}{'位辩手' if is_zh else ' debaters'})", expanded=False):
                    # Enable/disable toggle
                    dept_enabled = st.checkbox(
                        f"{t('config_dept_enabled')}/{t('config_dept_disabled')}",
                        value=True,
                        key=f"config_dept_enabled_{dept_key}",
                    )
                    
                    if dept_enabled:
                        # Debater list
                        for d_key, debater in debaters.items():
                            d_name = debater.get("zh_name", d_key) if is_zh else debater.get("en_name", d_key)
                            style_key = "zh_style" if is_zh else "en_style"
                            current_style = debater.get(style_key, "")
                            
                            st.markdown(f"**{d_name}**")
                            new_style = st.text_area(
                                f"{'提示词' if is_zh else 'Prompt'}",
                                value=current_style,
                                height=80,
                                key=f"config_debater_style_{dept_key}_{d_key}",
                                label_visibility="collapsed",
                            )
                            # Real-time update to config
                            if new_style != current_style:
                                config["departments"][dept_key]["debaters"][d_key][style_key] = new_style
                                st.session_state.workgroup_config = config
            
            st.divider()
            
            # Global parameter editing — conditional on mode
            if is_animation and not is_academic:
                # Animation mode: show all creative params
                st.markdown("### " + t("config_global_params"))
                
                vd = config.get("visual_directive", {})
                vd_zh = vd.get("zh", "") if isinstance(vd, dict) else ""
                new_vd_zh = st.text_area(
                    t("config_visual_directive") + " (中文)",
                    value=vd_zh,
                    height=100,
                    key="config_visual_directive_zh",
                )
                if new_vd_zh != vd_zh:
                    if not isinstance(config.get("visual_directive"), dict):
                        config["visual_directive"] = {"zh": "", "en": ""}
                    config["visual_directive"]["zh"] = new_vd_zh
                    st.session_state.workgroup_config = config
                
                neg = config.get("negative_prompts", "")
                new_neg = st.text_input(
                    t("config_negative_prompts"),
                    value=neg,
                    key="config_negative_prompts",
                )
                if new_neg != neg:
                    config["negative_prompts"] = new_neg
                    st.session_state.workgroup_config = config
            else:
                # Academic/generic mode: simplified global params
                st.markdown("### " + ("辩论参数" if is_zh else "Debate Parameters"))
            
            dr = config.get("debate_rounds", 3)
            new_dr = st.slider(
                t("config_debate_rounds"),
                min_value=1, max_value=10, value=dr,
                key="config_debate_rounds",
            )
            if new_dr != dr:
                config["debate_rounds"] = new_dr
                st.session_state.workgroup_config = config
            
            st.divider()
            
            # Skill injection area
            st.markdown("### " + t("config_skill_injection"))
            skill_md = st.text_area(
                t("config_skill_injection"),
                height=80,
                placeholder=t("config_skill_injection_hint"),
                key="config_skill_md",
            )
            if skill_md:
                # Select target department
                dept_keys = list(depts.keys())
                target_options = [t("config_skill_target_all")] + [
                    depts[k].get("zh_name", k) if is_zh else depts[k].get("en_name", k)
                    for k in dept_keys
                ]
                selected_target = st.selectbox(
                    t("config_skill_target"),
                    options=target_options,
                    key="config_skill_target",
                )
                if st.button("💉 " + ("注入" if is_zh else "Inject"), use_container_width=True):
                    if selected_target == target_options[0]:
                        target_depts = None  # 全部
                    else:
                        # Find corresponding dept_key
                        idx = target_options.index(selected_target) - 1
                        target_depts = [dept_keys[idx]]
                    
                    config = merge_skill_injection(config, skill_md, target_depts)
                    st.session_state.workgroup_config = config
                    st.success("✅ " + ("Skill已注入！" if is_zh else "Skill injected!"))
    
    # Bottom action bar
    st.divider()
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        # Save config
        save_name = st.text_input(
            t("config_save_name"),
            value=st.session_state.get("workgroup_name", ""),
            key="config_save_name",
        )
        if st.button(t("config_save"), use_container_width=True):
            config = (st.session_state.get("workgroup_config") or get_current_config())
            if save_name:
                save_profile(save_name, config)
                set_last_used(save_name)
                st.success(f"✅ {'已保存为' if is_zh else 'Saved as'}: {save_name}")
            else:
                st.error("请输入配置名称" if is_zh else "Please enter a config name")
    
    with col2:
        # Export config
        if st.button(t("config_export"), use_container_width=True):
            config = (st.session_state.get("workgroup_config") or get_current_config())
            json_str = export_config(config)
            st.code(json_str, language="json")
            st.download_button(
                "⬇️ " + ("下载JSON" if is_zh else "Download JSON"),
                data=json_str,
                file_name=_safe_filename(f"{st.session_state.get('workgroup_name', 'config')}.json"),
                mime="application/json",
            )
    
    with col3:
        # Import config
        imported_json = st.text_area(
            "JSON" if is_zh else "JSON",
            height=60,
            placeholder="粘贴JSON配置..." if is_zh else "Paste JSON config...",
            key="config_import_json",
            label_visibility="collapsed",
        )
        if st.button(t("config_import"), use_container_width=True):
            if imported_json:
                try:
                    cfg = import_config(imported_json)
                    errors = validate_config(cfg)
                    if errors:
                        for e in errors:
                            st.warning(e)
                    apply_config(cfg)
                    st.session_state.workgroup_config = cfg
                    st.success("✅ " + ("配置已导入" if is_zh else "Config imported"))
                except Exception as e:
                    st.error(str(e))
            else:
                st.warning("请先粘贴JSON" if is_zh else "Please paste JSON first")
    
    with col4:
        # Confirm and start
        if st.button(t("config_apply"), type="primary", use_container_width=True):
            config = (st.session_state.get("workgroup_config") or get_current_config()) or {}
            apply_config(config)
            cfg_name = config.get("name", "自定义配置")
            st.session_state.workgroup_name = cfg_name
            set_last_used(cfg_name)
            st.success(t("config_applied"))



# ============ v4.0 Requirement Research Tab ============

def render_requirement_tab():
    """Render requirement research Tab — Phase 0~4 complete flow"""
    is_zh = st.session_state.get("lang", "zh") == "zh"
    
    
    
    
    st.header("🔬 " + ("需求调研" if is_zh else "Requirement Research"))
    st.caption("v4.0 " + ("对话式需求调研 → 结构化 → 讨论组 → 配置推荐 → 审核" if is_zh else "Interview → Structure → Discussion → Recommend → Review"))
    
    # ---- LLM Config (from sidebar) ----
    api_url = st.session_state.get("api_url", "")
    api_key = st.session_state.get("api_key", "")
    model = st.session_state.get("model_name", "")
    
    def _llm_call(system_prompt: str, user_prompt: str) -> str:
        """Unified LLM call, reusing app API configuration"""
        import requests as _req
        headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": 0.3,
            "max_tokens": 4096,
        }
        try:
            resp = _req.post(api_url, headers=headers, json=payload, timeout=120)
            resp.raise_for_status()
            return resp.json()["choices"][0]["message"]["content"]
        except Exception as e:
            return f"[ERROR] {str(e)}"
    
    llm_available = bool(api_url and api_key and model)
    
    # ---- Phase state management ----
    if "req_phase" not in st.session_state:
        st.session_state.req_phase = 0
    if "req_document" not in st.session_state:
        st.session_state.req_document = None
    if "req_structured" not in st.session_state:
        st.session_state.req_structured = None
    if "req_discussion" not in st.session_state:
        st.session_state.req_discussion = None
    if "req_config" not in st.session_state:
        st.session_state.req_config = None
    if "req_interview_history" not in st.session_state:
        st.session_state.req_interview_history = []
    
    phase = st.session_state.req_phase
    
    # ---- Progress bar ----
    phase_names = ["Phase 0: 需求调研", "Phase 1: 需求结构化", "Phase 2: 讨论组", "Phase 3: 配置推荐", "Phase 4: 用户审核"] if is_zh else \
                  ["Phase 0: Interview", "Phase 1: Structure", "Phase 2: Discussion", "Phase 3: Recommend", "Phase 4: Review"]
    progress_val = (phase + 1) / 5
    st.progress(progress_val)
    
    cols = st.columns(5)
    for i, name in enumerate(phase_names):
        with cols[i]:
            label = f"**{name}**" if i <= phase else name
            color = "✅" if i < phase else ("🔵" if i == phase else "⬜")
            st.markdown(f"{color} {label}")
    
    st.divider()
    
    # ================================================================
    # Phase 0: Conversational requirement interview
    # ================================================================
    if phase == 0:
        st.subheader("Phase 0: " + ("对话式需求调研" if is_zh else "Requirement Interview"))
        
        # Domain selection
        domain_options = {
            "academic_research": "📚 学术调研" if is_zh else "📚 Academic Research",
            "animation": "🎬 动画制作" if is_zh else "🎬 Animation",
            "programming_tutorial": "💻 程序与教程" if is_zh else "💻 Programming & Tutorial",
            "general": "🌐 通用" if is_zh else "🌐 General",
        }
        selected_domain = st.selectbox(
            "选择领域" if is_zh else "Domain",
            options=list(domain_options.keys()),
            format_func=lambda x: domain_options[x],
            key="req_domain_select",
        )
        
        # Initial requirement input
        user_topic = st.text_area(
            "描述你的需求" if is_zh else "Describe your requirement",
            placeholder="例如：我想调研机器学习在能源经济学中的前沿应用..." if is_zh else "e.g., I want to research ML applications in energy economics...",
            height=120,
            key="req_topic_input",
        )
        
        # Existing interview history
        if st.session_state.req_interview_history:
            st.markdown("#### " + ("访谈记录" if is_zh else "Interview History"))
            for i, turn in enumerate(st.session_state.req_interview_history):
                with st.chat_message("user" if turn["role"] == "user" else "assistant"):
                    st.markdown(turn["content"])
        
        # Follow-up question input
        follow_up = st.chat_input(
            ("回答追问 / 补充信息" if st.session_state.req_interview_history else "开始调研") if is_zh else \
            ("Answer follow-up / Add info" if st.session_state.req_interview_history else "Start research"),
            key="req_follow_up_input",
        )
        
        if follow_up:
            if not llm_available:
                st.error("⚠️ " + ("请先在侧边栏配置LLM API" if is_zh else "Please configure LLM API in sidebar"))
            else:
                # Add user message
                st.session_state.req_interview_history.append({"role": "user", "content": follow_up})
                
                # AI follow-up
                with st.spinner("🤔 " + ("调研AI正在分析..." if is_zh else "Research AI analyzing...")):
                    # Initialize or continue interview
                    if "req_interviewer" not in st.session_state or st.session_state.get("req_interviewer") is None:
                        st.session_state.req_interviewer = RequirementInterviewer(llm_call_fn=_llm_call, domain_hint=selected_domain, language="zh" if is_zh else "en")
                        result = st.session_state.req_interviewer.start(user_input=follow_up)
                    else:
                        result = st.session_state.req_interviewer.chat(user_message=follow_up)
                    
                    # Save AI follow-up
                    next_q = result.get("question", "")
                    if next_q:
                        understanding = result.get("current_understanding", "")
                        ai_msg = ""
                        if understanding:
                            ai_msg = "📋 " + ("当前理解" if is_zh else "Current understanding") + ":\n" + understanding + "\n\n---\n"
                        ai_msg += f"**{("追问" if is_zh else "Follow-up")}:** {next_q}"
                        st.session_state.req_interview_history.append({"role": "assistant", "content": ai_msg})
                    
                    # Save requirement doc (if interview complete)
                    if result.get("is_complete", False):
                        st.session_state.req_document = st.session_state.req_interviewer.get_requirement_document()
                
                st.rerun()
        
        # End interview button
        if st.session_state.req_interview_history:
            col1, col2 = st.columns(2)
            with col1:
                if st.button("✅ " + ("需求已明确，进入结构化" if is_zh else "Requirement clear, proceed to structure"), type="primary", key="req_finish_interview"):
                    # If no auto-generated doc, manually build from user input
                    if not st.session_state.req_document:
                        st.session_state.req_document = RequirementDocument(
                            topic=user_topic or follow_up or "",
                            domain=selected_domain,
                            objectives=[],
                            constraints={},
                            key_questions=[],
                            deliverable_type="",
                            quality_criteria="",
                            interview_history=st.session_state.req_interview_history,
                            domain_specific={},
                        )
                    st.session_state.req_phase = 1
                    st.rerun()
            with col2:
                if st.button("🔄 " + ("重新开始" if is_zh else "Restart"), key="req_restart_interview"):
                    st.session_state.req_phase = 0
                    st.session_state.req_interview_history = []
                    st.session_state.req_document = None
                    st.rerun()
    
    # ================================================================
    # Phase 1: Requirement structuring
    # ================================================================
    elif phase == 1:
        st.subheader("Phase 1: " + ("需求结构化" if is_zh else "Requirement Structuring"))
        
        if st.session_state.req_document:
            req_doc = st.session_state.req_document
            
            # Display requirement doc
            with st.expander("📋 " + ("需求文档" if is_zh else "Requirement Document"), expanded=True):
                if isinstance(req_doc, dict):
                    st.json(req_doc)
                else:
                    st.json(req_doc.to_dict() if hasattr(req_doc, 'to_dict') else str(req_doc))
            
            # Structure button
            if st.button("🔧 " + ("结构化需求" if is_zh else "Structure Requirement"), type="primary", key="req_structure_btn"):
                with st.spinner("🤔 " + ("结构化中..." if is_zh else "Structuring...")):
                    structurer = RequirementStructurer(llm_call_fn=_llm_call if llm_available else None)
                    structured = structurer.structure(doc=req_doc)
                    st.session_state.req_structured = structured
                    st.session_state.req_phase = 2
                    st.rerun()
        else:
            st.warning("⚠️ " + ("请先完成Phase 0的需求调研" if is_zh else "Please complete Phase 0 interview first"))
            if st.button("⬅️ " + ("返回Phase 0" if is_zh else "Back to Phase 0"), key="req_back_p0"):
                st.session_state.req_phase = 0
                st.rerun()
    
    # ================================================================
    # Phase 2: Discussion group
    # ================================================================
    elif phase == 2:
        st.subheader("Phase 2: " + ("需求讨论组" if is_zh else "Requirement Discussion Group"))
        
        if st.session_state.req_structured:
            structured = st.session_state.req_structured
            
            # Display structured results
            with st.expander("📊 " + ("结构化需求" if is_zh else "Structured Requirement"), expanded=False):
                if isinstance(structured, dict):
                    st.json(structured)
                else:
                    st.json(structured.to_dict() if hasattr(structured, 'to_dict') else str(structured))
            
            # Discussion group config preview
            st.markdown("#### " + ("讨论组成员" if is_zh else "Discussion Panel"))
            roles = structured.get("suggested_roles", []) if isinstance(structured, dict) else (structured.suggested_roles if hasattr(structured, 'suggested_roles') else [])
            if roles:
                for role_info in roles:
                    role_name = role_info.get("role", "") if isinstance(role_info, dict) else str(role_info)
                    reason = role_info.get("reason", "") if isinstance(role_info, dict) else ""
                    st.markdown(f"- **{role_name}**: {reason}")
            else:
                st.info(("将使用默认3角色：方法论审查 / 反方视角 / 落地可行性" if is_zh else "Default 3 roles: Methodology Review / Counter-perspective / Feasibility"))
            
            # Run discussion
            if st.button("💬 " + ("运行讨论组" if is_zh else "Run Discussion"), type="primary", key="req_discuss_btn"):
                with st.spinner("🤔 " + ("讨论组正在审议..." if is_zh else "Discussion panel deliberating...")):
                    discussion = DiscussionGroup(llm_call_fn=_llm_call if llm_available else None)
                    result = discussion.discuss(requirement=structured)
                    st.session_state.req_discussion = result
                    st.session_state.req_phase = 3
                    st.rerun()
        else:
            st.warning("⚠️ " + ("请先完成Phase 1的结构化" if is_zh else "Please complete Phase 1 structuring first"))
            if st.button("⬅️ " + ("返回Phase 1" if is_zh else "Back to Phase 1"), key="req_back_p1"):
                st.session_state.req_phase = 1
                st.rerun()
    
    # ================================================================
    # Phase 3: Config recommendation
    # ================================================================
    elif phase == 3:
        st.subheader("Phase 3: " + ("部门配置推荐" if is_zh else "Department Config Recommendation"))
        
        if st.session_state.req_discussion and st.session_state.req_structured:
            discussion = st.session_state.req_discussion
            structured = st.session_state.req_structured
            
            # Display discussion results
            with st.expander("💬 " + ("讨论纪要" if is_zh else "Discussion Minutes"), expanded=False):
                if isinstance(discussion, dict):
                    st.json(discussion)
                else:
                    st.json(discussion.to_dict() if hasattr(discussion, 'to_dict') else str(discussion))
            
            # Generate config recommendation
            if st.button("⚙️ " + ("生成部门配置" if is_zh else "Generate Department Config"), type="primary", key="req_config_btn"):
                with st.spinner("🤔 " + ("正在推荐部门配置..." if is_zh else "Generating config recommendation...")):
                    recommender = ConfigRecommender(llm_call_fn=_llm_call if llm_available else None)
                    config = recommender.recommend(requirement=structured, discussion=discussion)
                    st.session_state.req_config = config
                    st.session_state.req_phase = 4
                    st.rerun()
        else:
            st.warning("⚠️ " + ("请先完成Phase 2的讨论组" if is_zh else "Please complete Phase 2 discussion first"))
            if st.button("⬅️ " + ("返回Phase 2" if is_zh else "Back to Phase 2"), key="req_back_p2"):
                st.session_state.req_phase = 2
                st.rerun()
    
    # ================================================================
    # Phase 4: User review
    # ================================================================
    elif phase == 4:
        st.subheader("Phase 4: " + ("用户审核" if is_zh else "User Review"))
        
        if st.session_state.req_config:
            config = st.session_state.req_config
            
            st.markdown("### " + ("推荐的部门配置" if is_zh else "Recommended Department Config"))
            
            # Full presentation + collapsible
            departments = config.get("departments", {}) if isinstance(config, dict) else (config.get("departments", {}) if hasattr(config, 'get') else {})
            dept_order = config.get("dept_order", list(departments.keys())) if isinstance(config, dict) else []
            
            # Basic info editing
            config_name = st.text_input(
                ("配置名称" if is_zh else "Config Name"),
                value=config.get("name", "") if isinstance(config, dict) else "",
                key="req_config_name",
            )
            config_desc = st.text_area(
                ("配置描述" if is_zh else "Description"),
                value=config.get("description", "") if isinstance(config, dict) else "",
                key="req_config_desc",
            )
            
            # Department-by-department display (collapsible)
            edited_departments = {}
            for dept_key in dept_order:
                dept = departments.get(dept_key, {})
                zh_name = dept.get("zh_name", dept_key) if isinstance(dept, dict) else dept_key
                en_name = dept.get("en_name", "") if isinstance(dept, dict) else ""
                
                with st.expander(f"🏢 {zh_name} / {en_name} (`{dept_key}`)", expanded=False):
                    # Department info editable
                    d_zh = st.text_input("Chinese name", value=zh_name, key=f"req_dept_zh_{dept_key}")
                    d_en = st.text_input("English name", value=en_name, key=f"req_dept_en_{dept_key}")
                    
                    debaters = dept.get("debaters", {})
                    if not isinstance(debaters, dict):
                        debaters = {} if isinstance(dept, dict) else {}
                    # Track deletions in session_state
                    del_key = f"_req_del_{dept_key}"
                    if del_key not in st.session_state:
                        st.session_state[del_key] = []
                    active_debaters = {k: v for k, v in debaters.items() if k not in st.session_state[del_key] and isinstance(v, dict)}
                    edited_debaters = {}
                    
                    for debater_key, debater in active_debaters.items():
                        db_col1, db_col2 = st.columns([6, 1])
                        with db_col1:
                            st.markdown(f"**Debater {debater_key}**: {debater.get('zh_name', '')}")
                        with db_col2:
                            can_delete = len(active_debaters) > 2
                            if st.button("🗑️", key=f"req_db_del_{dept_key}_{debater_key}", disabled=not can_delete, help=("至少保留2位辩手" if can_delete else "至少保留2位辩手") if not can_delete else ("删除辩手" if is_zh else "Remove debater")):
                                st.session_state[del_key].append(debater_key)
                                st.rerun()
                        
                        db_zh_name = st.text_input("Debater Chinese name", value=debater.get("zh_name", ""), key=f"req_db_zh_{dept_key}_{debater_key}")
                        db_en_name = st.text_input("Debater English name", value=debater.get("en_name", ""), key=f"req_db_en_{dept_key}_{debater_key}")
                        db_zh_style = st.text_area("Chinese style (Prompt)", value=debater.get("zh_style", ""), height=100, key=f"req_db_zhs_{dept_key}_{debater_key}")
                        db_en_style = st.text_area("English style (Prompt)", value=debater.get("en_style", ""), height=80, key=f"req_db_ens_{dept_key}_{debater_key}")
                        
                        edited_debaters[debater_key] = {
                            "zh_name": db_zh_name,
                            "en_name": db_en_name,
                            "zh_style": db_zh_style,
                            "en_style": db_en_style,
                        }
                    
                    # Add debater button
                    if st.button("➕ " + ("添加辩手" if is_zh else "Add Debater"), key=f"req_db_add_{dept_key}", use_container_width=True):
                        existing_keys = set(active_debaters.keys())
                        next_key = None
                        for c in "ABCDEFGHIJKLMNOPQRSTUVWXYZ":
                            if c not in existing_keys and c not in st.session_state[del_key]:
                                next_key = c
                                break
                        if next_key:
                            new_debater = {
                                "zh_name": f"新辩手{next_key}" if is_zh else f"New Debater {next_key}",
                                "en_name": f"New Debater {next_key}",
                                "zh_style": "请补充辩手风格描述" if is_zh else "Describe debater style",
                                "en_style": "Describe debater style",
                            }
                            if isinstance(st.session_state.req_config, dict):
                                depts = st.session_state.req_config.get("departments", {})
                                if dept_key in depts and isinstance(depts[dept_key], dict):
                                    depts[dept_key].setdefault("debaters", {})[next_key] = new_debater
                            st.rerun()
                    
                    edited_departments[dept_key] = {
                        "zh_name": d_zh,
                        "en_name": d_en,
                        "debaters": edited_debaters,
                    }

            
            # Review operations
            st.divider()
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("✅ " + ("确认配置，进入辩论" if is_zh else "Confirm & Start Debate"), type="primary", key="req_confirm_btn"):
                    # Write reviewed config to session_state for smart grouping tab
                    final_config = {
                        "name": config_name,
                        "description": config_desc,
                        "departments": edited_departments,
                        "dept_order": list(edited_departments.keys()),
                        "p2_cross_debates": config.get("p2_cross_debates", []) if isinstance(config, dict) else [],
                        "p5_cross_debates": config.get("p5_cross_debates", []) if isinstance(config, dict) else [],
                        "proofread_departments": config.get("proofread_departments", []) if isinstance(config, dict) else [],
                        "debate_rounds": config.get("debate_rounds", 2) if isinstance(config, dict) else 2,
                        "negative_prompts": config.get("negative_prompts", "") if isinstance(config, dict) else "",
                    }
                    
                    # Write to current config
                    apply_config(final_config)
                    st.session_state.workgroup_config = final_config
                    
                    # Set pipeline_mode explicitly based on department keys
                    _p4_dept_keys = set(final_config.get("departments", {}).keys())
                    _academic_keys = {"literature_search", "methodology_review", "report_integration", 
                                     "programming", "tutorial", "metadata_inspector", "citation_network",
                                     "data_validation", "counter_evidence", "topic_clustering", "visualization"}
                    st.session_state.pipeline_mode = "academic" if _p4_dept_keys & _academic_keys else "animation"
                    
                    # Set auto-jump flag for toast notification
                    st.session_state._auto_jump_to_debate = True
                    
                    # Auto-fill requirement doc into script field so debate tab can start directly
                    req_doc = st.session_state.get("req_document")
                    req_structured = st.session_state.get("req_structured")
                    script_parts = []
                    
                    # Extract from req_document (supports both object and dict)
                    if req_doc:
                        # Get attributes from object or dict
                        def _get(obj, key, default=None):
                            if isinstance(obj, dict):
                                return obj.get(key, default)
                            return getattr(obj, key, default) if hasattr(obj, key) else default
                        
                        _topic = _get(req_doc, "topic")
                        if _topic:
                            script_parts.append(f"研究主题：{_topic}")
                        _objectives = _get(req_doc, "objectives") or []
                        if _objectives:
                            script_parts.append(f"研究目标：{'；'.join(str(o) for o in _objectives)}")
                        _constraints = _get(req_doc, "constraints") or {}
                        if isinstance(_constraints, dict):
                            for k, v in _constraints.items():
                                script_parts.append(f"{k}：{v}")
                        _key_questions = _get(req_doc, "key_questions") or []
                        if _key_questions:
                            script_parts.append(f"关键问题：{'；'.join(str(q) for q in _key_questions)}")
                        _domain_specific = _get(req_doc, "domain_specific") or {}
                        if isinstance(_domain_specific, dict):
                            for k, v in _domain_specific.items():
                                script_parts.append(f"{k}：{v}")
                        # Fallback: raw_text if no structured fields
                        _raw = _get(req_doc, "raw_text")
                        if not script_parts and _raw:
                            script_parts.append(_raw[:500])
                    
                    # Extract from req_structured
                    if req_structured:
                        _hints = None
                        if isinstance(req_structured, dict):
                            _hints = req_structured.get("department_hints")
                        elif hasattr(req_structured, "department_hints"):
                            _hints = req_structured.department_hints
                        if _hints:
                            hints = [h["description"] for h in _hints if isinstance(h, dict) and "description" in h]
                            if hints:
                                script_parts.append(f"部门方向：{'；'.join(hints)}")
                    
                    # Fallback: use interview history if script_parts still empty
                    if not script_parts:
                        _history = st.session_state.get("req_interview_history", [])
                        for turn in _history:
                            if turn.get("role") == "user" and turn.get("content"):
                                script_parts.append(turn["content"])
                                if len(script_parts) >= 3:
                                    break
                    
                    # Final fallback: use req_topic_input
                    if not script_parts:
                        _topic_input = st.session_state.get("req_topic_input", "")
                        if _topic_input:
                            script_parts.append(_topic_input)
                    
                    # Always set script content (even if empty, to clear stale state)
                    _fill_content = "\n".join(script_parts) if script_parts else ""
                    if _fill_content:
                        st.session_state.script = _fill_content
                        st.session_state.script_input = _fill_content
                        st.session_state.config_user_input = _fill_content
                        # Set pending fill flag for render_input_tab to pick up
                        st.session_state._pending_script_fill = _fill_content
                    
                    st.session_state.req_phase = 0  # 重置
                    st.session_state.req_interview_history = []
                    st.session_state.req_document = None
                    st.session_state.req_structured = None
                    st.session_state.req_discussion = None
                    st.session_state.req_config = None
                    
                    # Auto-start debate: set flag for main() to pick up after rerun
                    st.session_state._auto_start_debate = True
                    st.success("✅ " + ("配置已确认，正在自动跳转到辩论..." if is_zh else "Config confirmed, auto-starting debate..."))
                    st.balloons()
                    st.rerun()
            
            with col2:
                if st.button("🔄 " + ("修改配置" if is_zh else "Edit Config"), key="req_edit_btn"):
                    # Stay in Phase 4 for user to continue editing
                    pass
            
            with col3:
                if st.button("⬅️ " + ("退回Phase 3" if is_zh else "Back to Phase 3"), key="req_back_p3"):
                    st.session_state.req_phase = 3
                    st.rerun()
            
            # Save config as preset
            with st.expander("💾 " + ("导出为预设" if is_zh else "Export as Preset"), expanded=False):
                preset_name = st.text_input(("预设名称" if is_zh else "Preset Name"), value=config_name, key="req_export_name")
                if st.button(("保存预设" if is_zh else "Save Preset"), key="req_save_preset_btn"):
                    final_config = {
                        "name": preset_name,
                        "description": config_desc,
                        "departments": edited_departments,
                        "dept_order": list(edited_departments.keys()),
                        "p2_cross_debates": config.get("p2_cross_debates", []) if isinstance(config, dict) else [],
                        "p5_cross_debates": config.get("p5_cross_debates", []) if isinstance(config, dict) else [],
                        "proofread_departments": config.get("proofread_departments", []) if isinstance(config, dict) else [],
                        "debate_rounds": config.get("debate_rounds", 2) if isinstance(config, dict) else 2,
                        "negative_prompts": config.get("negative_prompts", "") if isinstance(config, dict) else "",
                    }
                    save_profile(preset_name, final_config)
                    st.success(("预设已保存！" if is_zh else "Preset saved!"))

            

        else:
            st.warning("⚠️ " + ("请先完成Phase 3的配置推荐" if is_zh else "Please complete Phase 3 recommendation first"))
            if st.button("⬅️ " + ("返回Phase 3" if is_zh else "Back to Phase 3"), key="req_back_p3_empty"):
                st.session_state.req_phase = 3
                st.rerun()



def main():
    st.set_page_config(
        page_title="AI Consensus Pipeline" if st.session_state.get("lang", "zh") == "zh" else "AI Consensus Pipeline",
        page_icon="🧠",
        layout="wide",
    )
    

    # ---- Keepalive: prevent Streamlit Cloud idle sleep (15min timeout) ----
    # Sends a HEAD request every 10 minutes to keep container awake.
    # Invisible to user, does not trigger reruns or state changes.
    st_components.html(
        """
        <script>
        (function(){
            var interval = 10 * 60 * 1000;  // 10 minutes
            setInterval(function(){
                fetch(window.location.origin + window.location.pathname, {method:'HEAD', mode:'no-cors', cache:'no-store'});
            }, interval);
        })();
        </script>
        """,
        height=0,
    )

    # Language welcome page (first visit)
    if not st.session_state.get("_lang_selected", False):
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown("<h2 style=\"text-align: center;\">🌐 Choose Language / 选择语言</h2>", unsafe_allow_html=True)
            st.markdown("<br>", unsafe_allow_html=True)
            _welcome_lang = st.radio(
                "",
                options=["zh", "en"],
                format_func=lambda x: "🇨🇳 中文" if x == "zh" else "🇬🇧 English",
                horizontal=True,
                key="_welcome_lang_radio"
            )
            if st.button("▶️ 进入 / Enter", type="primary", use_container_width=True, key="_welcome_enter"):
                st.session_state.lang = _welcome_lang
                st.session_state._lang_selected = True
                st.rerun()
        st.stop()

    render_sidebar()
    
    st.title(t("title"))
    st.caption(t("subtitle"))
    
    # Auto-jump to debate tab after Phase 4 confirmation
    if st.session_state.get("_auto_start_debate"):
        st.session_state._auto_start_debate = False
        st.session_state._tab_index = 1  # debate tab
        # Clear previous results before starting new debate
        st.session_state.debate_completed = False
        st.session_state.dept_results = {}
        st.session_state.final_output = {}
        st.session_state._debate_running = True  # auto-start debate
    
    # Tab index tracked by _tab_index (not widget-bound, safe to set directly)
    if "_tab_index" not in st.session_state:
        st.session_state._tab_index = 0
    
    _debating = st.session_state.get("_debate_running", False)
    
    # If debate is running, force tab to debate tab
    if _debating:
        st.session_state._tab_index = 1
    
    _tab_labels = [
        t("tab_setup"),
        t("tab_debate_combined"),
        t("tab_output_combined"),
        t("tab_tools"),
    ]
    # If our desired tab differs from radio's current value, force radio reinit
    _desired_tab = st.session_state._tab_index
    if st.session_state.get("_main_tab_radio") != _desired_tab:
        st.session_state.pop("_main_tab_radio", None)

    def _on_tab_change():
        """Sync radio click back to _tab_index"""
        st.session_state._tab_index = st.session_state._main_tab_radio

    _active_tab = st.radio(
        "",
        options=range(len(_tab_labels)),
        format_func=lambda i: _tab_labels[i],
        horizontal=True,
        index=_desired_tab,
        key="_main_tab_radio",
        on_change=_on_tab_change,
        label_visibility="collapsed",
    )
    
    # Tab0: 需求与配置 — sub-tabs
    with st.container():
        if _active_tab == 0:
            sub_tab0, sub_tab1, sub_tab2 = st.tabs([
                t("tab_setup_req"),
                t("tab_setup_config"),
                t("tab_setup_input"),
            ])
            with sub_tab0:
                render_requirement_tab()
            with sub_tab1:
                render_config_tab()
            with sub_tab2:
                render_input_tab()
    
    # Tab1: 辩论 — expanders
    if _active_tab == 1:
        with st.expander(t("tab_dept_debate"), expanded=True):
            render_debate_tab()
        with st.expander(t("tab_cross_debate_sub"), expanded=False):
            render_cross_tab()
    
    # Tab2: 产出 — expanders
    if _active_tab == 2:
        with st.expander(t("tab_final_output"), expanded=True):
            render_output_tab()
        with st.expander(t("tab_proofread_sub"), expanded=False):
            render_proofread_tab()
    
    # Tab3: 工具 — expanders
    if _active_tab == 3:
        with st.expander(t("tab_compare_sub"), expanded=True):
            render_compare_tab()
        with st.expander(t("tab_market_sub"), expanded=False):
            render_market_tab()
    
    # Auto-jump toast after tab rendering
    if st.session_state.get("_auto_jump_to_debate"):
        st.session_state._auto_jump_to_debate = False
        is_zh = st.session_state.get("lang", "zh") == "zh"
        st.toast(t("toast_config_confirmed"))


if __name__ == "__main__":
    init_state()
    main()


